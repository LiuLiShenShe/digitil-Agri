#!/usr/bin/env python3
"""
Generate paper from 计算机研究与发展专题投稿初稿.md
Exactly matching the author's draft content.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from xml.sax.saxutils import escape
import os
import re

BASE = '/data/fj/数字孪生-paper-work'
FIGURES_DIR = os.path.join(BASE, '论文/figures/智能体数字孪生图')
OUTPUT = os.path.join(BASE, '论文/KAFarmTwin-知识增强多智能体数字孪生场景构建.docx')

PAGE_WIDTH = Emu(7560310)
PAGE_HEIGHT = Emu(10692130)
TOP_MARGIN = Emu(648335)
BOTTOM_MARGIN = Emu(540385)
LEFT_MARGIN = Emu(612140)
RIGHT_MARGIN = Emu(612140)
HEADER_DISTANCE = Emu(720000)
FOOTER_DISTANCE = Emu(720000)
ONE_COLUMN_SPACE = 720
TWO_COLUMN_SPACE = 425
IN_TWO_COLUMN_BODY = False

# Figure file mapping
FIGURE_FILES = {
    1: '图 1：总体框架图.png',
    2: '图 2：对象本体与关系图.png',
    3: '图 3：多智能体协作流程图.png',
    4: '图 4：多保真资产路由图.png',
    5: '图 5：对象级长期记忆图.png',
    6: '图 6：规则校验与轨迹图.png',
    7: '图7_系统原型界面多面板图_使用用户截图.png',
    8: '图 8 ：主实验结构可靠性对比图.png',
    9: '图 9 不同消融版本的结构可靠性对比.png',
}

REFERENCE_ORDER = [
    1, 2, 4, 5, 6, 7, 22, 23, 24, 8, 9, 10, 11, 12, 13, 14,
    28, 29, 30, 31, 33, 15, 18, 25, 26, 27, 19, 21, 34, 37,
]
CITATION_MAP = {old: new for new, old in enumerate(REFERENCE_ORDER, start=1)}

def _compact_citations(numbers):
    numbers = sorted(dict.fromkeys(numbers))
    ranges = []
    start = prev = None
    for n in numbers:
        if start is None:
            start = prev = n
        elif n == prev + 1:
            prev = n
        else:
            ranges.append((start, prev))
            start = prev = n
    if start is not None:
        ranges.append((start, prev))
    parts = [str(a) if a == b else f'{a}-{b}' for a, b in ranges]
    return ','.join(parts)

def remap_citations(text):
    """Map old reference numbers to first-appearance order in manuscript text."""
    def repl(match):
        content = match.group(1)
        nums = []
        for part in content.split(','):
            part = part.strip()
            if '-' in part:
                a, b = [int(x.strip()) for x in part.split('-', 1)]
                nums.extend(range(a, b + 1))
            else:
                nums.append(int(part))
        mapped = [CITATION_MAP[n] for n in nums]
        return f'[{_compact_citations(mapped)}]'
    return re.sub(r'\[([0-9]+(?:\s*[-,]\s*[0-9]+)*)\]', repl, text)

def set_section_columns(section, num_columns=1, space=ONE_COLUMN_SPACE):
    """Set Word section column count with explicit OOXML, matching the journal template."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc_grid = sect_pr.find(qn('w:docGrid'))
        if doc_grid is not None:
            sect_pr.insert(sect_pr.index(doc_grid), cols)
        else:
            sect_pr.append(cols)

    cols.set(qn('w:space'), str(space))
    if num_columns > 1:
        cols.set(qn('w:num'), str(num_columns))
    else:
        cols.attrib.pop(qn('w:num'), None)

def configure_section(section, num_columns=1, column_space=None):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    set_section_columns(
        section,
        num_columns,
        column_space if column_space is not None
        else (TWO_COLUMN_SPACE if num_columns > 1 else ONE_COLUMN_SPACE),
    )

def switch_columns(doc, num_columns):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(section, num_columns=num_columns)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    return section

def start_two_column_body(doc):
    global IN_TWO_COLUMN_BODY
    switch_columns(doc, 2)
    IN_TWO_COLUMN_BODY = True

def begin_full_width_block(doc):
    if IN_TWO_COLUMN_BODY:
        switch_columns(doc, 1)

def end_full_width_block(doc):
    if IN_TWO_COLUMN_BODY:
        switch_columns(doc, 2)

# ============================================================
# Helper functions
# ============================================================

def add_para(doc, text, font_name='宋体', font_size=Pt(10.5),
             bold=False, alignment=None, first_line_indent=None,
             space_before=0, space_after=0, line_spacing=1.15,
             font_name_west='Times New Roman'):
    """Add a paragraph with specified formatting."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_name_west:
            run._element.rPr.rFonts.set(qn('w:ascii'), font_name_west)
            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name_west)
        run.font.size = font_size
        run.bold = bold
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    return p

def add_mixed_para(doc, segments, alignment=None, first_line_indent=None,
                   space_before=0, space_after=0, line_spacing=1.15):
    """segments: list of (text, font_name, font_size, bold, font_name_west)"""
    p = doc.add_paragraph()
    for seg in segments:
        text = seg[0]
        fn = seg[1] if len(seg) > 1 else '宋体'
        fs = seg[2] if len(seg) > 2 else Pt(10.5)
        b = seg[3] if len(seg) > 3 else False
        fnw = seg[4] if len(seg) > 4 else 'Times New Roman'
        run = p.add_run(text)
        run.font.name = fn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        run._element.rPr.rFonts.set(qn('w:ascii'), fnw)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), fnw)
        run.font.size = fs
        run.bold = b
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    return p

def add_figure(doc, fig_num, en_caption, cn_caption, width=Inches(4.25)):
    """Add figure with bilingual caption."""
    begin_full_width_block(doc)
    fname = FIGURE_FILES.get(fig_num)
    if fname:
        fpath = os.path.join(FIGURES_DIR, fname)
        if os.path.exists(fpath):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(2)
            p_img.paragraph_format.space_after = Pt(1)
            run = p_img.add_run()
            run.add_picture(fpath, width=width)

    add_para(doc, f'Fig. {fig_num} {en_caption}',
             font_name='Times New Roman', font_size=Pt(8),
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    add_para(doc, f'图{fig_num} {cn_caption}',
             font_name='宋体', font_size=Pt(8),
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    end_full_width_block(doc)

def add_section_heading(doc, number, title):
    """一级标题: 小四黑 (12pt)"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    rn = p.add_run(f'{number} ')
    rn.font.name = 'Times New Roman'
    rn.font.size = Pt(12)
    rn.bold = True
    rt = p.add_run(title)
    rt.font.name = '黑体'
    rt._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    rt._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    rt._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rt.font.size = Pt(12)
    rt.bold = True
    return p

def add_subsection_heading(doc, number, title):
    """二级标题: 五号黑 (10.5pt)"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15
    rn = p.add_run(f'{number} ')
    rn.font.name = 'Times New Roman'
    rn.font.size = Pt(10.5)
    rn.bold = True
    rt = p.add_run(title)
    rt.font.name = '黑体'
    rt._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    rt._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    rt._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rt.font.size = Pt(10.5)
    rt.bold = True
    return p

def add_body(doc, text, indent=Emu(266700)):
    """Body text: 五号宋体"""
    text = remap_citations(text)
    return add_para(doc, text, font_name='宋体', font_size=Pt(10),
                    first_line_indent=indent, line_spacing=1.15)

def add_ref(doc, text):
    """Reference: 小五号"""
    return add_para(doc, text, font_name='Times New Roman', font_size=Pt(9),
                    line_spacing=1.05)

def add_equation(doc, latex_str):
    """Add a math equation using OMML (Office Math Markup Language).
    This is a simplified version - for complex equations we use plain text with Unicode."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.first_line_indent = Emu(266700)

    # Use italic Times New Roman for math
    run = p.add_run(latex_str)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_display_math(doc, text):
    """Display math equation as an Office math object, centered."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    math_xml = f'''<m:oMathPara {nsdecls("m", "w")}>
        <m:oMath>
            <m:r>
                <m:rPr><m:nor/></m:rPr>
                <w:rPr>
                    <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:eastAsia="Cambria Math"/>
                    <w:color w:val="000000"/>
                    <w:sz w:val="20"/>
                    <w:szCs w:val="20"/>
                </w:rPr>
                <m:t>{escape(text)}</m:t>
            </m:r>
        </m:oMath>
    </m:oMathPara>'''
    p._p.append(parse_xml(math_xml))
    return p

def make_three_line_table(doc, headers, data, en_caption, cn_caption, note=None, bold_rows=None):
    """Three-line table per journal format."""
    begin_full_width_block(doc)
    add_para(doc, en_caption, font_name='Times New Roman', font_size=Pt(8), bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=1, space_after=0)
    add_para(doc, cn_caption, font_name='宋体', font_size=Pt(8), bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1)

    n_rows = len(data) + 1
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = True

    # Three-line borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(6.8)
        run.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # bot border under header
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:tcBorders>')
        tcPr.append(tcB)

    if bold_rows is None:
        bold_rows = []

    # Data
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            run.font.size = Pt(6.8)
            if j == 0:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            else:
                run.font.name = 'Times New Roman'
            if i + 1 in bold_rows:
                run.bold = True

    if note:
        add_para(doc, note, font_name='宋体', font_size=Pt(8), space_before=3, space_after=6)

    end_full_width_block(doc)
    return table


# ============================================================
# BUILD DOCUMENT
# ============================================================

doc = Document()
section = doc.sections[0]
configure_section(section, num_columns=1)

header = section.header
header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
header_p.text = ''
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = header_p.add_run('计算机研究与发展    Journal of Computer Research and Development    卷(期):起止页,年')
hr.font.name = '宋体'
hr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
hr._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
hr._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
hr.font.size = Pt(9)

footer = section.footer
footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_p.text = ''
footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
fr = footer_p.add_run('收稿日期、修回日期、基金项目和通信作者信息隐去，供双盲评审。')
fr.font.name = '宋体'
fr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
fr._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
fr._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
fr.font.size = Pt(7.5)

# ============================================================
# CHINESE TITLE PAGE
# ============================================================

cn_title = '面向设施农业数字孪生的知识增强多智能体协作与可追溯场景构建方法'
add_para(doc, cn_title, font_name='宋体', font_size=Pt(16), bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6,
         font_name_west='Times New Roman')

add_para(doc, '（作者信息隐去，供双盲评审）',
         font_name='宋体', font_size=Pt(14),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)

add_para(doc, '（单位信息隐去，供双盲评审）',
         font_name='宋体', font_size=Pt(9),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)

add_para(doc, '（通信作者邮箱隐去）',
         font_name='宋体', font_size=Pt(9),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

# ============================================================
# ENGLISH TITLE
# ============================================================

en_title = 'Knowledge-Augmented Multi-Agent Collaboration for Traceable Scene Construction in Protected-Agriculture Digital Twins'
add_para(doc, en_title,
         font_name='Times New Roman', font_size=Pt(14), bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=3)

add_para(doc, '(Author information withheld for double-blind review)',
         font_name='Times New Roman', font_size=Pt(10.5),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)

add_para(doc, '(Affiliation information withheld for double-blind review)',
         font_name='Times New Roman', font_size=Pt(9),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

abstract_cn = (
    '设施农业数字孪生场景构建需要保证农业对象层级、三维资产、运行数据和生产事件一致绑定。'
    '现有系统依赖人工建模与配置，通用大模型智能体虽能生成候选场景，但易出现层级缺失、绑定错误、'
    '资产不匹配和过程不可审计。本文提出KAFarmTwin知识约束多智能体方法，将对象本体、对象级长期记忆、'
    '多保真资产知识、规则校验和执行轨迹注入规划、布局、资产路由、数据绑定与校验闭环。30条任务评测'
    '表明，KAFarmTwin的关系F1、绑定F1、规则冲突率和可执行轨迹可信度分别为0.803、0.775、0.007和'
    '1.000；对象F1为0.711，体现保守的可验证对象生成策略。消融结果表明，本体、记忆、资产路由和规则'
    '校验分别支撑层级约束、历史证据、资产选择和冲突收敛。'
)

# ============================================================
# ENGLISH ABSTRACT
# ============================================================

add_para(doc, 'Abstract', font_name='Times New Roman', font_size=Pt(10.5), bold=True,
         space_before=6, space_after=3)

abstract_en = (
    'Protected-agriculture digital twins require consistent links among agricultural-object hierarchies, '
    '3D assets, runtime data and production events. Existing systems rely heavily on manual modelling '
    'and configuration. Although LLM agents can generate candidate scenes from natural-language requests, '
    'they often miss intermediate objects, produce invalid bindings, select mismatched assets or leave '
    'unauditable execution traces. This paper presents KAFarmTwin, a knowledge-constrained multi-agent '
    'method that injects an agricultural object ontology, object-level long-term memory, multi-fidelity '
    'asset knowledge, rule validation and executable traces into planning, layout generation, asset '
    'routing, data binding and validation. On 30 protected-agriculture tasks under a fair baseline '
    'protocol, KAFarmTwin achieves 0.803 Relation-F1, 0.775 Binding-F1, 0.007 rule violation rate and '
    '1.000 executable-trace faithfulness. Its Object-F1 is 0.711, indicating a conservative strategy '
    'that favours verifiable objects over unconstrained expansion. Ablations show that ontology, memory, '
    'asset routing and validation support hierarchy constraints, historical evidence, asset selection '
    'and conflict convergence, respectively.'
)
add_body(doc, abstract_en)

add_para(doc, '', font_name='Times New Roman', font_size=Pt(10.5), space_before=3, space_after=0)
add_mixed_para(doc, [
    ('Keywords ', 'Times New Roman', Pt(10.5), True),
    ('protected agriculture; digital twin; knowledge-augmented AI; multi-agent system; '
     'object ontology; long-term memory; rule validation; traceable construction',
     'Times New Roman', Pt(10.5), False)
], space_before=0, space_after=6)

# ============================================================
# CHINESE ABSTRACT
# ============================================================

add_para(doc, '摘  要', font_name='黑体', font_size=Pt(10.5), bold=True,
         space_before=6, space_after=3)

add_para(doc, abstract_cn, font_name='楷体', font_size=Pt(10.5),
         first_line_indent=Emu(266700), line_spacing=1.5)

add_mixed_para(doc, [
    ('关键词 ', '黑体', Pt(10.5), True),
    ('设施农业；数字孪生；知识增强人工智能；多智能体；对象本体；长期记忆；规则校验；可追溯构建',
     '楷体', Pt(10.5), False)
], space_before=3, space_after=3)

add_para(doc, '中图法分类号 TP391.9；S126', font_name='宋体', font_size=Pt(10.5),
         space_before=3, space_after=6)

start_two_column_body(doc)

# ============================================================
# 1. INTRODUCTION
# ============================================================

add_section_heading(doc, '1', '引言')

add_body(doc,
    '数字孪生通过虚拟模型、物理对象和运行数据之间的动态映射，为复杂系统的状态感知、过程分析'
    '和决策支持提供了重要技术路径[1,2,4]。在设施农业中，温室生产管理正在从环境级监测走向对象级、'
    '过程化和可追溯管理[5-7,22-24]。传统监控系统通常以温度、湿度、光照和设备状态等指标为中心，'
    '能够描述局部生产环境，却难以统一表示温室、地块、作物行、单株、传感器、摄像头、灌溉设备、'
    '表型指标和生产事件之间的空间、语义和时序关系。三维数字孪生为上述对象提供了共同的空间载体，'
    '使传感器数据、表型记录、灌溉事件和历史状态能够绑定到具体农业对象，从而支撑对象级查询、'
    '异常定位和生产过程复盘。')

add_body(doc,
    '然而，现有农业数字孪生和三维可视化系统主要强调监测、展示和数据汇聚[5,8,22]，场景构建仍'
    '高度依赖人工建模、人工拖拽和人工配置。开发者需要从自然语言或业务需求中识别农业对象，手工'
    '建立温室、地块、作物行和植株之间的层级关系，再为传感器、摄像头和设备配置数据接口、资产模型'
    '和空间位置。该过程效率低、复用性差，并且容易造成对象关系断裂。例如，同一传感器在三维场景中'
    '存在模型对象，在业务系统中存在设备对象，在时序数据库中又存在数据源对象。若缺少统一对象图和'
    '绑定规则，后续问答或决策系统很难判断数据究竟归属于哪个温室、作物行或植株。')

add_body(doc,
    '大模型智能体为数字孪生场景构建提供了新的自动化可能。智能体可以理解"构建一个30m×8m的番茄'
    '温室，包含作物行、传感器、摄像头和滴灌设备"等自然语言需求，并通过工具调用生成对象列表、'
    '三维布局和数据绑定[9-14,28-31,33]。但是，直接使用通用大模型或普通智能体仍面临3个关键挑战：'
    '1）对象图结构难以保证，模型可能漏掉地块、作物行等中间层级，或生成方向错误的包含（contains）、'
    '监测（monitors）和控制（controls）关系；2）三维资产、业务对象和运行数据难以一致绑定，例如将'
    '水泵资产绑定为植株模型，或生成缺少单位、时间戳和对象归属的表型指标；3）执行过程缺少可复核'
    '证据，声明式执行轨迹可以描述"做了校验"，但无法证明系统实际调用了哪些工具、使用了哪些知识'
    '以及是否通过规则检查。')

add_body(doc,
    '针对上述挑战，本文提出KAFarmTwin，一种面向设施农业数字孪生的知识增强多智能体协作与可追溯'
    '场景构建方法。其基本思想是：大模型负责语义理解和候选生成，农业对象本体、对象级记忆、多保真'
    '资产知识和规则约束负责结构化约束、数据绑定和结果校验；多智能体协作负责将候选对象图转化为'
    '可加载、可绑定、可校验和可审计的数字孪生场景。本文主要贡献包括3个方面：')

contributions = [
    '1. 提出设施农业数字孪生对象图表示与规则体系，统一描述温室、地块、作物行、植株、传感器、'
    '摄像头、设备、表型指标、生产事件和三维资产之间的层级、绑定和校验约束。',
    '2. 提出知识约束的多智能体场景构建流程，使农业对象本体、对象级记忆、资产元数据和规则库进入'
    '场景规划、空间布局、资产路由、数据绑定和规则校验器闭环，并通过执行式轨迹记录可审计证据。',
    '3. 构建设施农业对象图评测协议和30条任务集，使用对象、关系和绑定的精确率/召回率/F1值，'
    '规则冲突率，以及声明式轨迹和执行式轨迹的双层可追溯指标评估场景构建可靠性。实验结果表明，'
    'KAFarmTwin在关系F1值、绑定F1值、规则冲突率和可执行轨迹可信度上分别达到0.803、0.775、'
    '0.007和1.000，但对象展开仍偏保守，对象F1值为0.711。'
]
for c in contributions:
    add_body(doc, c)

# ============================================================
# 2. RELATED WORK
# ============================================================

add_section_heading(doc, '2', '相关工作')

add_subsection_heading(doc, '2.1', '农业数字孪生与三维场景构建')

add_body(doc,
    '数字孪生研究从制造领域逐步扩展到农业生产过程建模和智能农场管理[1,2,4,5,22-23]。在农业场景中，'
    '相关工作通常关注传感器数据汇聚、作物状态监测、农机或温室设备管理以及三维可视化展示[5-8,24]。'
    '这些研究证明了数字孪生在农业状态感知中的价值，但多数系统仍以数据看板或可视化场景为核心，'
    '场景中的对象、业务实体和历史事件之间缺少可计算的语义关系。对于"温室-地块-作物行-植株-传感器-'
    '事件"这类多层对象关系，若三维模型仅作为展示元素存在，就难以支持对象级查询、历史追踪和规则'
    '校验。因此，本文关注的不是三维可视化质量本身，而是如何自动构建包含对象层级、数据绑定、资产'
    '来源和规则校验结果的数字孪生对象图。')

add_subsection_heading(doc, '2.2', '大模型智能体与工具调用')

add_body(doc,
    '大模型智能体通过结合语言理解、任务分解、工具调用和结果汇总，使模型从单轮文本生成扩展到'
    '交互式任务执行[9-14,28-31,33]。推理-行动框架（ReAct）强调推理与行动的交替组织[11]，工具学习'
    '框架（Toolformer）表明语言模型可以学习调用外部工具[12]，链式思维和树状搜索进一步'
    '增强了大模型的问题分解和候选探索能力[31,33]，多智能体研究则进一步讨论了角色分工、任务调度和'
    '协作执行[13-14,30]。对于数字孪生场景构建，智能体可以调用模型检索、布局求解、对象查询、数据'
    '绑定和规则校验等工具，将自然语言需求转化为可执行配置。问题在于，通用智能体通常缺少设施农业'
    '对象层级和业务规则，可能生成表面完整但结构不合法的场景。仅增加智能体数量也不能自然保证正确性，'
    '错误对象和错误绑定可能在多个智能体之间继续传播。因此，多智能体协作需要与可执行知识约束和可回流'
    '校验机制结合。')

add_subsection_heading(doc, '2.3', '知识增强与神经符号融合')

add_body(doc,
    '知识增强人工智能通过检索增强生成（RAG）、知识图谱、本体、规则系统和外部记忆等方式，将符号'
    '知识引入神经模型推理过程[9-10,15,18]。RAG方法能够缓解模型知识不足问题，但检索到的文档并不'
    '自动成为可执行约束[9-10]。知识图谱和本体能够显式表示概念、属性和关系，适合描述设施农业中的'
    '对象层级、设备控制关系和数据归属关系[15,25-27]。其中，农业领域本体库和传感器语义本体为'
    '农业概念、观测指标、传感器、执行器和采样过程的规范化表示提供了参考[25-27]。神经符号融合进一步'
    '强调，神经模型可负责感知、语言理解和候选生成，符号知识可负责结构约束、规则校验和错误修正[18]。'
    '本文将这一思想落实到数字孪生对象图构建任务中，使知识不只作为提示词或检索材料出现，而是进入规划、'
    '布局、资产路由、绑定和校验全过程。')

add_subsection_heading(doc, '2.4', '长期记忆与可追溯推理')

add_body(doc,
    '长期记忆和可解释性研究关注模型输出是否具有可追踪证据、可复核过程和明确的适用边界[13,19,28-29]。'
    '在设施农业中，植株长势、环境状态、灌溉事件、设备维护和异常告警都具有时间属性。如果系统只生成'
    '静态三维场景，就无法回答"第3行番茄最近7天长势如何"或"水泵最近24小时是否异常"等对象级问题。'
    '另一方面，可追溯推理要求记录智能体在每一步使用了哪些知识、调用了哪些工具、产生了哪些输出以及'
    '是否通过校验。本文将对象级记忆和智能体执行轨迹纳入数字孪生构建流程，使场景不仅可见，而且可查、'
    '可验、可追踪；同时在评测中区分声明式轨迹和带有证据编号或调用编号的执行式轨迹，避免将模型自述'
    '过程误认为真实执行证据。')

# ============================================================
# 3. METHOD
# ============================================================

add_section_heading(doc, '3', 'KAFarmTwin知识增强多智能体场景构建方法')

add_body(doc,
    '图1给出KAFarmTwin的总体框架。框架由需求输入、知识层、智能体层、数字孪生对象图和应用层构成：'
    '自然语言需求被解析为场景构建任务，对象本体、对象级记忆、资产知识和规则库约束规划、布局、'
    '资产路由、数据绑定和校验，最终输出可加载、可绑定、可校验的对象图。')

# FIG-1: Architecture
add_figure(doc, 1,
    'Overall architecture of KAFarmTwin framework',
    'KAFarmTwin总体框架图')

add_body(doc,
    'KAFarmTwin的设计围绕引言中提出的3类挑战展开。针对对象图结构不稳定问题，系统在规划阶段引入'
    '农业对象本体和R1-R3层级/空间规则；针对资产、业务和数据绑定不一致问题，系统在资产路由和数据'
    '绑定阶段引入多保真资产元数据、对象级记忆和R4-R6、R8-R9规则；针对过程不可复核问题，系统记录'
    '工具调用、策略状态和证据编号，并由R7和R10检查轨迹完整性与修正链条。')

add_subsection_heading(doc, '3.1', '问题定义')

add_body(doc,
    '本文将设施农业数字孪生场景构建定义为受知识约束的对象图生成任务。给定用户需求、已有数据、'
    '资产库、对象记忆和规则集合，系统需要生成可加载到三维场景、可映射到业务对象、可连接运行数据'
    '且可被规则复核的对象图。输入定义为')

add_display_math(doc, 'Q = {q, D_s, A, M_t, R},')

add_body(doc,
    '其中，q表示用户自然语言需求，D_s表示已有传感器数据、表型数据和生产事件，A表示三维资产库，'
    'M_t表示对象级长期记忆，R表示农业规则库。三维资产库包含已有GLB格式模型、F2DMAS高保真植株'
    '资产路径、TRELLIS.2三维快速生成任务、程序化模型和占位资产。输出定义为',
    indent=None)

add_display_math(doc, 'Y = {G, B, V, T},')

add_body(doc,
    '其中，G表示三维数字孪生场景图，B表示对象与资产、数据、事件和业务对象之间的绑定关系，'
    'V表示规则校验结果，T表示智能体执行轨迹。本文目标是在农业知识集合K的约束下生成综合评分较高'
    '的数字孪生对象图：',
    indent=None)

add_display_math(doc, 'Y* = arg max_Y S(Y | Q, K),')

add_body(doc,
    '其中K包含农业对象本体、资产知识、对象级记忆和规则约束。与直接生成场景结构化数据不同，'
    '本文要求结果同时满足对象层级、资产来源、数据绑定、规则校验和可追溯执行记录。')

add_subsection_heading(doc, '3.2', '农业对象知识表示')

add_body(doc,
    '本文将设施农业对象知识表示为')

add_display_math(doc, 'K_o = (C, R_o, P, I),')

add_body(doc,
    '其中，C表示对象类别集合，R_o表示对象关系集合，P表示属性集合，I表示对象实例集合。对象类别'
    '覆盖温室（Greenhouse）、地块（Plot）、作物行（CropRow）、植株（Plant）、传感器（Sensor）、'
    '摄像头（Camera）、设备（Device）、性状（Trait）、事件（Event）和资产（Asset）。典型关系包括'
    '包含（contains）、隶属（belongs_to）、监测（monitors）、观测（observes）、控制（controls）、'
    '关联资产（has_asset）、具有关联性状（has_trait）和关联事件（has_event）。',
    indent=None)

add_body(doc,
    '图2展示设施农业对象本体与关系。该本体为"番茄温室""作物行""传感器""灌溉设备"等候选对象提供'
    '可计算的类型边界、层级关系和绑定约束。')

# FIG-2: Ontology
add_figure(doc, 2,
    'Agricultural object ontology and relation model',
    '设施农业对象本体与关系图')

add_body(doc,
    '如表1所示，本文将KAFarmTwin的知识增强能力拆分为对象本体、对象级长期记忆、多保真资产路由、'
    '规则校验器和智能体执行轨迹5个模块，并将每个模块对应到后续实验指标。')

# TABLE 1: Core modules
t1_headers = ['模块', '主要知识', '作用', '对应实验指标']
t1_data = [
    ['农业对象本体', '对象类别、层级关系、语义关系', '约束温室、地块、作物行、植株和设备之间的对象图结构', '关系正确率（RA）、层级错误率'],
    ['对象级长期记忆', '时序指标、事件、日报、历史状态', '支持历史查询、数据绑定和对象状态追踪', '绑定准确率（BA）、轨迹完整率（TC）、R8'],
    ['多保真资产路由', '资产元数据、保真度、成本、缺失任务', '选择F2DMAS高保真资产、GLB格式模型、TRELLIS.2生成任务、程序化或占位资产', '资产路由准确率（AR）、规则冲突率（VR）'],
    ['规则校验器（Validator）', '对象规则、空间规则、绑定规则、资产规则', '发现并修正对象层级、空间布局和资产绑定冲突', '规则冲突率（VR）、规则校验器冲突率'],
    ['智能体执行轨迹（Agent Trace）', '智能体、工具、输入输出、状态、策略', '记录规划、布局、资产、绑定和校验全过程', '轨迹完整率（TC）'],
]
make_three_line_table(doc, t1_headers, t1_data,
    'Table 1 Core modules of KAFarmTwin',
    '表1 本文方法核心模块及作用')

add_body(doc,
    '由表1可见，各模块分别约束对象层级、历史证据、资产选择、规则收敛和执行留痕，形成后续多智能体'
    '协作与实验评价的共同基础。')

add_subsection_heading(doc, '3.3', '多智能体协作框架')

add_body(doc,
    'KAFarmTwin将场景构建流程拆分为编排、规划、布局、资产路由、数据绑定和校验6类智能体。各智能体'
    '分别完成任务分解、对象图生成、三维坐标计算、资产选择、对象-数据绑定和规则复核。')

add_body(doc,
    '图3给出多智能体协作流程。编排器将需求拆解为子任务，各智能体结果写入轨迹记录器；当校验器发现'
    '对象缺失、关系方向错误、资产不匹配或绑定缺失时，冲突被路由回相应智能体修正。')

# FIG-3: Multi-agent
add_figure(doc, 3,
    'Multi-agent collaboration workflow',
    '多智能体协作流程图')

add_body(doc,
    '该流程把大模型候选生成与符号约束闭环结合起来，使多智能体从角色拆分机制转化为"生成、校验、'
    '修正、留痕"的可复核流程。')

add_subsection_heading(doc, '3.4', '多保真资产路由机制')

add_body(doc,
    '设施农业数字孪生不要求所有对象使用同一精度模型。重点植株需要较高几何可信度，背景对象更关注'
    '加载效率；资产库缺失时应生成占位对象和补资产任务，而不是中断构建流程。')

add_body(doc,
    '本文将资产选择定义为',
    indent=None)

add_display_math(doc, 'a_i* = arg max_{a∈A} (α·S_f(o_i, a) + β·S_q(a) + γ·S_e(o_i) − λ·C(a)),')

add_body(doc,
    '其中，S_f表示对象与资产的语义匹配度，S_q表示资产质量评分，S_e表示对象在任务中的重要性，'
    'C(a)表示调用或生成成本，α, β, γ, λ为权重。资产保真度智能体根据该策略在F2DMAS高保真资产路径、'
    'TRELLIS.2三维快速生成任务、轻量GLB格式模型、程序化模型和占位任务之间做选择。')

add_body(doc,
    '图4展示多保真资产路由机制，系统按对象类型、任务重要性、资产质量和调用成本选择高保真、轻量、'
    '程序化或占位资产。')

# FIG-4: Asset Router
add_figure(doc, 4,
    'Multi-fidelity asset routing mechanism',
    '多保真资产路由机制图')

add_body(doc,
    '本文将F2DMAS作为重点植株的高保真资产路由接入路径：当对象被识别为重点植株、异常植株或论文样本'
    '时，资产保真度智能体生成F2DMAS路由决策并写入对象图；普通缺失设备或背景对象则生成TRELLIS.2类'
    '任务或占位模型[21,34,37]。因此，资产库不完备会转化为可追踪的补资产任务。')

add_subsection_heading(doc, '3.5', '对象级长期记忆与数据绑定')

add_body(doc,
    '设施农业数字孪生对象不仅是三维模型，也是状态记忆单元。传感器语义网和轻量化观测本体已表明，'
    '传感器、观测、采样和执行器可被统一建模为可查询的语义资源[26-27]。本文将对象o_i的记忆定义为')

add_display_math(doc, 'M(o_i) = {P_i, S_i^t, E_i^t, A_i, T_i},')

add_body(doc,
    '其中，P_i表示静态属性，如对象类型、位置和尺寸；S_i^t表示动态状态，如温度、湿度、光照、CO₂、'
    '株高和冠幅；E_i^t表示事件记录，如灌溉、施肥、病害、采样和告警；A_i表示关联三维资产；T_i表示'
    '智能体操作记录。数据绑定关系定义为',
    indent=None)

add_display_math(doc, 'B = {(o_i, d_j, r_{ij}, t)},')

add_body(doc,
    '表示对象o_i在时间t与数据d_j建立关系r_{ij}。通过这种绑定，系统能够围绕对象检索历史状态，'
    '而不是生成缺少对象来源的概括性回答。')

add_body(doc,
    '图5给出对象级长期记忆与动态数据绑定机制。每个农业对象关联静态属性、三维资产、时序数据、'
    '表型指标、生产事件和智能体操作记录，使历史查询具有明确对象边界。')

# FIG-5: Memory
add_figure(doc, 5,
    'Object-level long-term memory and dynamic data binding mechanism',
    '对象级长期记忆与动态数据绑定机制图')

add_subsection_heading(doc, '3.6', '规则校验与智能体执行轨迹')

add_body(doc,
    '本文规则库包含对象层级、数据绑定、空间布局、资产类型、摄像头、设备覆盖、执行轨迹、记忆查询、'
    '缺失资产和错误修正10类规则。如表2所示，这些检查点覆盖对象图构建、资产绑定和执行轨迹的关键约束。')

# TABLE 2: Rules
t2_headers = ['规则', '检查内容']
t2_data = [
    ['R1', '对象层级合法：温室包含地块或栽培区，地块包含作物行或苗床，作物行包含植株。'],
    ['R2', '数据绑定合法：传感器、表型和事件数据必须有绑定对象、单位和时间戳。'],
    ['R3', '空间布局合法：对象不悬空、不越界，作物行或苗床位于地块或温室边界内。'],
    ['R4', '资产类型一致：对象类型与GLB格式模型、F2DMAS高保真资产、TRELLIS.2三维生成任务、程序化或占位资产策略一致。'],
    ['R5', '摄像头合法：摄像头必须有位姿、观测目标和视场覆盖关系。'],
    ['R6', '设备覆盖合法：灌溉、水肥、补光和通风设备必须绑定控制区域或服务对象。'],
    ['R7', '智能体执行轨迹完整：至少记录规划、布局、资产路由、数据绑定和校验步骤。'],
    ['R8', '记忆查询合法：历史查询必须限制对象、指标、时间范围、事件类型和返回条数。'],
    ['R9', '缺失资产不中断：缺失GLB格式模型时必须生成占位对象和资产生成任务。'],
    ['R10', '错误可修正：规则冲突必须输出冲突类型、触发规则和修正方案。'],
]
make_three_line_table(doc, t2_headers, t2_data,
    'Table 2 Rule checkpoints R1-R10',
    '表2 规则检查点R1-R10')

add_body(doc,
    '由表2可见，R1-R6约束对象、空间、资产和设备绑定，R7-R10进一步约束执行轨迹、记忆查询、缺失资产'
    '和错误修正，使校验结果能够直接定位到可修复环节。')

add_body(doc,
    '规则冲突数量定义为',
    indent=None)

add_display_math(doc, 'V(Y) = Σ_{r_k∈R} I[r_k(Y) = false],')

add_body(doc,
    '目标是在保持对象语义完整的前提下降低V(Y)。智能体执行轨迹记录每一步的智能体、工具、输入输出、'
    '状态和耗时[28-29]；文件写入、任意HTTP请求、直接数据库写入和真实设备控制等高风险操作被禁止或'
    '仅允许预览。若关键步骤缺失、工具调用没有证据编号或冲突没有修正建议，输出不能被视为完整可审计结果。')

add_body(doc,
    '图6展示规则校验与轨迹记录流程，冲突类型、触发规则、修正建议和执行状态均写入轨迹，便于定位'
    '具体智能体和工具调用步骤。')

# FIG-6: Rule Validation + Trace
add_figure(doc, 6,
    'Rule validation and Agent Trace recording process',
    '规则校验与Agent Trace记录流程图')

# ============================================================
# 4. SYSTEM IMPLEMENTATION
# ============================================================

add_section_heading(doc, '4', '系统实现')

add_body(doc,
    '本文在智慧农业数字孪生原型系统中实现KAFarmTwin，用于验证知识约束多智能体流程能否落地为可调用'
    '的场景构建服务。前端负责三维场景、对象树、自然语言输入、验收控制台、Agent Trace和资产路由展示；'
    '后端提供农业对象管理、场景绑定、对象级记忆、资产治理、语义构建和验收聚合接口。')

add_body(doc,
    '后端采用对象管理、场景绑定、记忆管理、资产治理和Agent编排分层实现。对象服务维护Greenhouse、'
    'Plot、CropRow、Plant、Sensor、Device和Camera等业务对象；记忆服务提供时序、事件和日报查询；'
    '资产服务维护GLB、F2DMAS、TRELLIS.2、程序化和占位策略；Agent编排服务组织规划、布局、资产路由、'
    '数据绑定和规则校验，并记录各阶段执行证据。')

add_body(doc,
    '为支撑可重复验证，原型实现提供固定验收任务，输入需求为"搭建番茄温室，包含20株番茄、气象站、水泵、'
    '摄像头和传感器"。验收服务聚合语义构建、对象计数、资产路由、业务绑定、对象记忆、校验问题、日报源'
    '和归档准备状态。前端验收界面展示端到端阶段状态、对象数量、执行轨迹、资产路由、对象上下文、校验'
    '问题和温室日报摘要。')

add_body(doc,
    '图7展示系统原型界面，包括三维场景视图、验收控制台、Agent Trace面板和资产路由面板。四类界面分别'
    '对应本文方法的场景输出、端到端验收、可追溯执行记录和多保真资产选择证据。')

# FIG-7: System prototype
add_figure(doc, 7,
    'KAFarmTwin system prototype interface',
    'KAFarmTwin系统原型界面多面板图')

add_body(doc,
    '原型系统用于验证对象图构建、资产路由、数据绑定和Trace机制，不等同于完整生产级农业控制平台；'
    '真实设备闭环控制、长期运行稳定性和更大规模作物数据将在后续工作中扩展。')

# ============================================================
# 5. EXPERIMENTS
# ============================================================

add_section_heading(doc, '5', '实验与分析')

add_subsection_heading(doc, '5.1', '实验设置')

add_body(doc,
    '本文实验围绕3个问题展开：1）在共享输出结构、对象知识和规则文本的公平条件下，KAFarmTwin是否比'
    '直接生成或普通智能体更能生成合法对象图；2）农业对象本体、对象级记忆、资产路由和规则校验器分别'
    '贡献哪些能力；3）方法收益是否依赖单一底座模型。本文构建30条设施农业数字孪生任务，覆盖场景构建、'
    '资产路由、数据绑定、规则修正和历史查询5类，每类6条。如表3所示，各类任务分别对应不同规则集合和'
    '能力边界。')

# TABLE 3: Task categories
t3_headers = ['任务类别', '数量', '代表任务', '主要考察能力', '对应规则']
t3_data = [
    ['场景构建', '6', 'T01番茄温室；T03玉米表型观测区', '对象识别、层级关系、空间布局、基础绑定', 'R1、R2、R3、R5、R6、R7'],
    ['资产路由', '6', 'T07重点番茄F2DMAS；T10缺失虫情灯/AI摄像头', '多保真资产选择、占位模型、生成任务', 'R3、R4、R5、R6、R7、R9'],
    ['数据绑定', '6', 'T13环境传感器指标；T18温室日报数据源', '传感器、摄像头、表型、事件和日报绑定', 'R1、R2、R5、R6、R7、R8'],
    ['规则修正', '6', 'T20作物行越界；T24水泵资产绑定错误', '冲突检测、规则触发、自动修正', 'R1、R2、R3、R4、R5、R6、R7、R9、R10'],
    ['历史查询', '6', 'T25长势变化；T30今日生产日报', '对象记忆检索、时序/事件查询、可追溯回答', 'R1、R2、R4、R5、R6、R7、R8'],
]
make_three_line_table(doc, t3_headers, t3_data,
    'Table 3 Task categories in the experiment',
    '表3 实验任务类别构成')

add_body(doc,
    '由表3可见，任务集并不只考察场景对象生成，还覆盖资产缺失、数据绑定、冲突修正和历史查询，因而能够'
    '检验对象图在构建后是否可查询、可校验和可追溯。')

add_subsection_heading(doc, '5.2', '对比方法')

add_body(doc,
    '为避免"完整工程系统对比裸大模型"的不公平设置，本文采用公平基线实验。所有非本文方法均调用同一'
    '语言模型，并接收相同的输出结构、农业对象类型、关系谓词、资产类型和R1-R10规则文本；标准对象数、'
    '标准关系数和标准绑定数仅用于离线评分，不提供给模型。对比方法包括Direct-LLM + Schema、'
    'Ontology/Rules Prompt、RAG-Agent、Single-Agent + Validator和Multi-Agent + Shared Knowledge。'
    '这些基线分别检验结构约束、规则提示、检索、一次性校验和角色分工的作用；本文方法在相同底座上额外'
    '启用对象本体、记忆、资产路由、Validator闭环和执行式Trace。如表4所示，所有方法共享输入约束，差异'
    '仅体现在知识是否进入工具化闭环。')

# TABLE 4: Baselines
t4_headers = ['方法', '共享输入', '差异设置', '预期暴露的问题']
t4_data = [
    ['Direct-LLM + Schema', '统一JSON schema、对象类型、关系谓词、资产类型、R1-R10', '一次性直接生成，不允许检索、工具调用或Validator。', '检验只有结构约束时的大模型直接生成能力。'],
    ['LLM + Ontology/Rules Prompt', '同上', '将对象本体和规则作为提示词显式注入，但不执行工具。', '检验看到规则文本是否足以形成可靠对象图。'],
    ['RAG-Agent + Ontology/Rules', '同上', '单智能体可检索同一份对象本体、规则和资产说明。', '检验检索知识能否转化为结构化约束。'],
    ['Single-Agent + Validator', '同上', '单智能体生成后执行一次离线规则校验器检查，但不回流修正。', '检验一次性校验与闭环修正的差异。'],
    ['Multi-Agent + Shared Knowledge', '同上', '多智能体分工，均可读取共享知识，但无闭环修正。', '检验角色分工在缺少闭环约束时的上限。'],
    ['Ours KAFarmTwin', '同上', '启用对象本体、记忆、资产路由、Validator闭环和执行式Trace。', '检验知识约束工具链对对象图可靠性的贡献。'],
]
make_three_line_table(doc, t4_headers, t4_data,
    'Table 4 Definitions of comparison methods',
    '表4 对比方法定义')

add_body(doc,
    '由表4可见，基线并非完全无知识，而是逐步加入结构约束、规则提示、检索、单次校验和多智能体分工；'
    '因此主实验重点比较可执行知识闭环相对于“看到知识文本”的增益。')

add_subsection_heading(doc, '5.3', '评价指标与评分细则')

add_body(doc,
    '本文使用对象、关系和绑定三个层面的精确率（Precision）、召回率（Recall）和F1值作为主指标，'
    '同时报告规则冲突率（VR）、轨迹字段完整率（TFC）和可执行轨迹可信度（ETF）。对对象、关系或'
    '绑定集合X，精确率、召回率和F1值定义为：',
    indent=None)

add_display_math(doc, 'P_X = N_{correct,X} / N_{generated,X},  R_X = N_{correct,X} / N_{required,X},  F1_X = 2P_X·R_X / (P_X + R_X).')

add_body(doc,
    '其中N_{correct,X}由标准对象、标准关系和标准绑定进行结构化匹配得到。对象匹配要求对象类型和任务'
    '语义一致；关系匹配要求主体、谓词和客体均正确，contains、belongs_to、monitors、observes、'
    'controls和has_asset等谓词必须方向正确；绑定匹配要求主体、目标和绑定类型完整，且数据、资产、事件'
    '或业务对象归属正确。该指标同时惩罚乱生成和漏生成，避免仅用对象完整率OC时掩盖选择性输出问题。')

add_body(doc,
    'VR根据任务预设规则检查点计算违反规则数量占比，R1、R2、R3、R4和R7视为致命约束。轨迹进一步拆分'
    '为TFC和ETF：TFC检查输出是否覆盖规划、布局、资产路由、数据绑定和校验等步骤字段；ETF仅统计来自'
    '系统工具调用链、带有证据编号或调用编号的执行式轨迹。Direct-LLM或普通智能体生成的声明式轨迹可'
    '计入TFC，但不计入高ETF。自动评分程序保留SR、OC、RA、BA和TC作为辅助诊断指标，主文分析以'
    'Object-F1、Relation-F1、Binding-F1、VR、TFC和ETF为核心。在最终实验中，本文方法的轨迹步骤'
    '均携带可复核证据编号，因此ETF达到1.000；其余基线仍主要停留在声明式轨迹，ETF维持为0。')

add_subsection_heading(doc, '5.4', '主实验结果')

add_body(doc,
    '如表5所示，v2公平基线实验同时报告Object-F1、Relation-F1、Binding-F1、VR、TFC和ETF，用于比较'
    '不同方法在对象、关系、绑定、规则一致性和轨迹可信性上的表现。')

# TABLE 5: Main experiment
t5_headers = ['方法', 'Object-F1↑', 'Relation-F1↑', 'Binding-F1↑', 'VR↓', 'TFC↑', 'ETF↑']
t5_data = [
    ['Direct-LLM + Schema', '0.814', '0.696', '0.554', '0.117', '0.067', '0.000'],
    ['LLM + Ontology/Rules Prompt', '0.835', '0.725', '0.533', '0.119', '0.027', '0.000'],
    ['RAG-Agent + Ontology/Rules', '0.819', '0.745', '0.658', '0.077', '0.040', '0.000'],
    ['Single-Agent + Validator', '0.837', '0.723', '0.595', '0.027', '0.133', '0.000'],
    ['Multi-Agent + Shared Knowledge', '0.827', '0.727', '0.625', '0.053', '0.973', '0.000'],
    ['Ours KAFarmTwin', '0.711', '0.803', '0.775', '0.007', '1.000', '1.000'],
]
make_three_line_table(doc, t5_headers, t5_data,
    'Table 5 Fair-baseline experiment results',
    '表5 公平基线实验结果', bold_rows=[6])

add_body(doc,
    '由表5可见，KAFarmTwin在Relation-F1、Binding-F1、VR和ETF上均优于或显著区别于非本文方法，但'
    'Object-F1并非最高，说明本文方法更偏向保守生成可验证对象关系。')

add_body(doc,
    '图8进一步以结构可靠性视角对比不同方法，重点展示Relation-F1、Binding-F1、规则通过率1−VR和ETF。')

# FIG-8: Main experiment
add_figure(doc, 8,
    'Structure reliability comparison on main experiment',
    '主实验结构可靠性对比图')

add_body(doc,
    '进一步结合表5结果可知，KAFarmTwin的Relation-F1和Binding-F1分别达到0.803和0.775，高于最佳非本文'
    '方法的0.745和0.658；VR降至0.007，说明规则冲突减少；TFC和ETF均达到1.000，说明轨迹字段完整且'
    '可由工具调用证据复核。相比之下，Multi-Agent + Shared Knowledge虽然TFC达到0.973，但ETF仍为0，'
    '表明多智能体可以生成较完整的过程描述，却不能证明这些过程来自真实执行链。')

add_body(doc,
    '同时，本文方法的Object-F1为0.711，低于若干直接生成式基线。该结果说明本文方法的优势不在于生成'
    '更多对象，而在于提高已生成对象之间的关系、绑定和规则一致性。换言之，知识约束闭环会抑制缺少资产'
    '证据、布局证据或绑定证据的对象扩张。该现象在农业数字孪生对象图构建中具有实际意义：对于后续查询'
    '和管理，错误绑定和错误层级往往比少量背景对象缺失更难修复。')

add_subsection_heading(doc, '5.5', '消融实验')

add_body(doc,
    '消融实验关注单个模块关闭后的机制性退化。为减少模型调用随机性影响，本文基于完整方法的结构化'
    '输出进行模块禁用消融，并重新计算关系正确率、资产路由准确率、层级错误率和Validator冲突率等'
    '模块级诊断指标。如表6所示，各消融版本仅关闭一个知识增强模块，以观察对应能力退化。')

# TABLE 6: Ablation
t6_headers = ['版本', 'OC↑', 'RA↑', 'AR↑', 'VR↓', 'TC↑', '层级错误率↓', 'Validator冲突率↓']
t6_data = [
    ['Ours', '0.524', '0.815', '0.597', '0.007', '0.993', '0.000', '0.008'],
    ['Ours w/o Ontology', '0.524', '0.473', '0.597', '0.108', '0.793', '1.000', '0.133'],
    ['Ours w/o Memory', '0.524', '0.721', '0.571', '0.162', '0.793', '0.000', '0.186'],
    ['Ours w/o Asset Router', '0.524', '0.731', '0.000', '0.108', '0.793', '0.000', '0.136'],
    ['Ours w/o Validator', '0.524', '0.815', '0.597', '0.628', '0.800', '0.154', '0.775'],
]
make_three_line_table(doc, t6_headers, t6_data,
    'Table 6 Ablation study results on knowledge enhancement modules',
    '表6 知识增强模块消融实验结果', bold_rows=[1])

add_body(doc,
    '由表6可见，去除本体、记忆、资产路由器或Validator会分别造成关系、资产路由、规则冲突和轨迹指标'
    '的退化，说明这些模块并非装饰性提示，而是参与了对象图收敛。')

add_body(doc,
    '图9从结构可靠性的角度展示各模块的独立贡献，便于直接比较完整方法与消融版本。')

# FIG-9: Ablation
add_figure(doc, 9,
    'Structure reliability comparison across ablation variants',
    '不同消融版本的结构可靠性对比图')

add_body(doc,
    '结合表6可以进一步看到，农业对象本体、对象级记忆、多保真资产路由和规则校验分别影响不同层面的可靠性。去除本体后，'
    'RA从0.815降至0.473，层级错误率升至1.000，说明对象本体主要负责温室、地块、作物行、植株和设备之间'
    '的结构约束。去除记忆后，RA、VR和TC同步变化，说明长期记忆主要支撑历史状态、事件证据和数据绑定的一致性。'
    '去除资产路由器后，AR降为0，说明资产路由直接决定高保真、轻量化、程序化和占位资产的选择。去除Validator'
    '后，VR从0.007升至0.628，Validator冲突率升至0.775，说明规则校验是冲突收敛和错误闭环的关键环节。')

add_subsection_heading(doc, '5.6', '错误归因与典型案例分析')

add_body(doc,
    '错误归因进一步解释了对象召回与结构可靠性之间的取舍。本文方法的低Object-F1主要来自保守式对象展开：'
    '缺失对象为309个，多余对象仅12个，且层级错误、资产类型错误和布局越界均为0；不可审计轨迹数量也降为0。'
    '典型任务显示，T01番茄温室可生成完整层级、绑定和执行链，T07可将重点植株路由至F2DMAS高保真资产路径，'
    'T24可纠正水泵误绑植株资产，T30可围绕对象记忆汇总生产日报。该结果说明，KAFarmTwin以对象召回为代价'
    '换取关系、绑定和轨迹合法性，后续优化应优先提升可验证对象召回，而不是放松规则约束。')

add_subsection_heading(doc, '5.7', '多模型配对鲁棒性补充')

add_body(doc,
    '为进一步证明KAFarmTwin的增益并非只来自某一个底座模型，补充实验采用配对式设计：对每个底座模型分别'
    '运行Base(M)和Ours(M)两组。Base(M)统一使用Direct-LLM + Schema口径，Ours(M)启用完整KAFarmTwin工具链。'
    '所有模型只替换模型接入层，不改变输出结构、评分脚本与方法流程。如表7所示，四个底座模型均完成配对评测。')

add_body(doc,
    '由表7可见，Ours(M)在四个底座上都提高了Relation-F1、Binding-F1与ETF，并将VR控制在0.007以内；与此同时，'
    'Object-F1并未总是同步上升，说明KAFarmTwin的优势主要体现在对象关系与工具闭环的可靠性，而不是盲目扩大对象数目。')

# TABLE 7: Multi-model robustness
t7_headers = ['底座模型', '设置', 'Object-F1↑', 'Relation-F1↑', 'Binding-F1↑', 'VR↓', 'TFC↑', 'ETF↑']
t7_data = [
    ['DeepSeek-V4-Flash', 'Base(M)', '0.841', '0.683', '0.485', '0.082', '0.167', '0.000'],
    ['DeepSeek-V4-Flash', 'Ours(M)', '0.612', '0.774', '0.716', '0.006', '0.987', '0.987'],
    ['GLM-5.1', 'Base(M)', '0.751', '0.537', '0.551', '0.065', '0.000', '0.000'],
    ['GLM-5.1', 'Ours(M)', '0.661', '0.805', '0.761', '0.007', '1.000', '1.000'],
    ['Kimi-K2.6', 'Base(M)', '0.851', '0.749', '0.597', '0.026', '0.033', '0.000'],
    ['Kimi-K2.6', 'Ours(M)', '0.680', '0.804', '0.794', '0.000', '1.000', '1.000'],
    ['MiniMax-M2.5', 'Base(M)', '0.870', '0.765', '0.606', '0.025', '0.000', '0.000'],
    ['MiniMax-M2.5', 'Ours(M)', '0.670', '0.819', '0.741', '0.007', '0.993', '0.993'],
]
make_three_line_table(doc, t7_headers, t7_data,
    'Table 7 Paired robustness results across base models',
    '表7 多模型配对鲁棒性补充结果')

add_body(doc,
    '表7进一步表明，KAFarmTwin的收益不依赖单一底座模型；不同底座上的提升模式一致，主要集中在关系、'
    '绑定和执行式轨迹可信性。')

# ============================================================
# 6. CONCLUSION
# ============================================================

add_section_heading(doc, '6', '结论')

add_body(doc,
    '本文提出KAFarmTwin，一种面向设施农业数字孪生对象图构建的知识约束多智能体框架。该方法将农业对象本体、'
    '对象级长期记忆、多保真资产知识、规则校验和智能体执行轨迹注入智能体规划、布局、资产路由、数据绑定和'
    '验证全过程，使大模型输出从对象罗列转向可验证、可追溯的数字孪生对象图。本文进一步将评测协议从单纯对象'
    '数量和轨迹字段完整性扩展为对象、关系和绑定的精确率/召回率/F1值、规则冲突率，以及声明式轨迹与执行式'
    '轨迹的双层可追溯指标。公平基线结果表明，在统一模型和统一知识输入下，KAFarmTwin在关系正确性、绑定'
    '有效性、规则一致性和执行式轨迹可信性上具有优势。消融实验说明，农业对象本体、对象级记忆、多保真资产路由和规则校验器分别对应层级关系、历史证据、'
    '资产选择和规则收敛等关键环节。')

add_body(doc,
    '本文仍存在3点局限。第一，当前方法对对象展开较保守，Object-F1低于若干直接生成式基线，后续需要在保持'
    '规则一致性的前提下提升对象召回和对象图覆盖范围。第二，原型验证主要围绕番茄温室和设施农业典型任务展开，'
    '仍需在更多作物类型、设备类型和生产流程上扩展任务集。第三，本文重点验证对象图构建、数据绑定、资产路由'
    '和轨迹机制，尚未覆盖真实设备闭环控制和长期生产运行验证。后续工作将进一步完善对象级记忆、资产治理和'
    '验收控制台之间的联动，并引入更大规模的真实农业数据流进行持续评估。')

# ============================================================
# AUTHOR CONTRIBUTIONS (placeholder for blind review)
# ============================================================

add_para(doc, '', font_name='宋体', font_size=Pt(10.5), space_before=6, space_after=3)
add_para(doc, '作者贡献声明：（作者信息隐去，供双盲评审）',
         font_name='宋体', font_size=Pt(10.5), bold=True,
         first_line_indent=Emu(266700), space_before=3, space_after=6)

# ============================================================
# REFERENCES
# ============================================================

add_para(doc, '', font_name='宋体', font_size=Pt(10.5), space_before=6, space_after=3)
add_para(doc, '参 考 文 献', font_name='黑体', font_size=Pt(12), bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

refs = [
    '[1] Grieves M. Digital Twin: Manufacturing Excellence through Virtual Factory Replication[R]. 2014.',
    '[2] Tao F, Zhang H, Liu A, Nee A Y C. Digital Twin in Industry: State-of-the-Art[J]. IEEE Transactions on Industrial Informatics, 2019, 15(4): 2405-2415.',
    '[3] Tao F, Sui F, Liu A, et al. Digital twin-driven product design, manufacturing and service with big data[J]. The International Journal of Advanced Manufacturing Technology, 2018, 94(9-12): 3563-3576.',
    '[4] Jones D, Snider C, Nassehi A, Yon J, Hicks B. Characterising the Digital Twin: A systematic literature review[J]. CIRP Journal of Manufacturing Science and Technology, 2020, 29: 36-52.',
    '[5] Pylianidis C, Osinga S, Athanasiadis I N. Introducing digital twins to agriculture[J]. Computers and Electronics in Agriculture, 2021, 184: 105942.',
    '[6] Wolfert S, Ge L, Verdouw C, Bogaardt M J. Big Data in Smart Farming: A review[J]. Agricultural Systems, 2017, 153: 69-80.',
    '[7] Liakos K G, Busato P, Moshou D, Pearson S, Bochtis D. Machine Learning in Agriculture: A Review[J]. Sensors, 2018, 18(8): 2674.',
    '[8] Drury B, Fernandes R, Moura M F, de Andrade Lopes A. A survey of semantic web technology for agriculture[J]. Information Processing in Agriculture, 2019, 6(4): 487-501.',
    '[9] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems 33. 2020: 9459-9474.',
    '[10] Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[EB/OL]. arXiv:2312.10997, 2023.',
    '[11] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]//International Conference on Learning Representations. 2023.',
    '[12] Schick T, Dwivedi-Yu J, Dessi R, et al. Toolformer: Language Models Can Teach Themselves to Use Tools[C]//Advances in Neural Information Processing Systems 36. 2023: 68539-68551.',
    '[13] Wang L, Ma C, Feng X, et al. A Survey on Large Language Model based Autonomous Agents[EB/OL]. arXiv:2308.11432, 2023.',
    '[14] Xi Z, Chen W, Guo X, et al. The Rise and Potential of Large Language Model Based Agents: A Survey[EB/OL]. arXiv:2309.07864, 2023.',
    '[15] Hogan A, Blomqvist E, Cochez M, et al. Knowledge Graphs[J]. ACM Computing Surveys, 2021, 54(4): 1-37.',
    '[16] Staab S, Studer R. Handbook on Ontologies[M]. Berlin: Springer, 2009.',
    '[17] Berners-Lee T, Hendler J, Lassila O. The Semantic Web[J]. Scientific American, 2001, 284(5): 34-43.',
    '[18] d\'Avila Garcez A, Lamb L C. Neurosymbolic AI: The 3rd wave[J]. Artificial Intelligence Review, 2023, 56(11): 12387-12406.',
    '[19] Arrieta A B, Diaz-Rodriguez N, Del Ser J, et al. Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI[J]. Information Fusion, 2020, 58: 82-115.',
    '[20] Amodei D, Olah C, Steinhardt J, et al. Concrete Problems in AI Safety[EB/OL]. arXiv:1606.06565, 2016.',
    '[21] Xiang J, Lv Z, Xu S, et al. Structured 3D Latents for Scalable and Versatile 3D Generation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2025.',
    '[22] Verdouw C, Tekinerdogan B, Beulens A, Wolfert S. Digital twins in smart farming[J]. Agricultural Systems, 2021, 189: 103046.',
    '[23] Walter A, Finger R, Huber R, Buchmann N. Smart farming is key to developing sustainable agriculture[J]. Proceedings of the National Academy of Sciences, 2017, 114(24): 6148-6150.',
    '[24] Kamilaris A, Prenafeta-Boldú F X. Deep learning in agriculture: A survey[J]. Computers and Electronics in Agriculture, 2018, 147: 70-90.',
    '[25] Jonquet C, Toulet A, Arnaud E, et al. AgroPortal: A vocabulary and ontology repository for agronomy[J]. Computers and Electronics in Agriculture, 2018, 144: 126-143.',
    '[26] Compton M, Barnaghi P, Bermudez L, et al. The SSN ontology of the W3C semantic sensor network incubator group[J]. Journal of Web Semantics, 2012, 17: 25-32.',
    '[27] Janowicz K, Haller A, Cox S J D, Le Phuoc D, Lefrançois M. SOSA: A lightweight ontology for sensors, observations, samples, and actuators[J]. Journal of Web Semantics, 2019, 56: 1-10.',
    '[28] Park J S, O\'Brien J, Cai C J, Morris M R, Liang P, Bernstein M S. Generative Agents: Interactive Simulacra of Human Behavior[C]//Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology. 2023: 1-22.',
    '[29] Shinn N, Cassano F, Gopinath A, Narasimhan K, Yao S. Reflexion: Language Agents with Verbal Reinforcement Learning[C]//Advances in Neural Information Processing Systems 36. 2023: 8634-8652.',
    '[30] Li G, Hammoud H, Itani H, Khizbullin D, Ghanem B. CAMEL: Communicative Agents for "Mind" Exploration of Large Language Model Society[C]//Advances in Neural Information Processing Systems 36. 2023: 51991-52008.',
    '[31] Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models[C]//Advances in Neural Information Processing Systems 35. 2022: 24824-24837.',
    '[32] Kojima T, Gu S S, Reid M, Matsuo Y, Iwasawa Y. Large Language Models Are Zero-Shot Reasoners[C]//Advances in Neural Information Processing Systems 35. 2022: 22199-22213.',
    '[33] Yao S, Yu D, Zhao J, et al. Tree of Thoughts: Deliberate Problem Solving with Large Language Models[C]//Advances in Neural Information Processing Systems 36. 2023: 11809-11822.',
    '[34] Lin C H, Gao J, Tang L, et al. Magic3D: High-Resolution Text-to-3D Content Creation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2023: 300-309.',
    '[35] Höllein L, Cao A, Owens A, Johnson J, Nießner M. Text2Room: Extracting Textured 3D Meshes from 2D Text-to-Image Models[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2023: 7875-7886.',
    '[36] Kerbl B, Kopanas G, Leimkühler T, Drettakis G. 3D Gaussian Splatting for Real-Time Radiance Field Rendering[J]. ACM Transactions on Graphics, 2023, 42(4): 139.',
    '[37] Yang Y, Sun F Y, Weihs L, et al. Holodeck: Language Guided Generation of 3D Embodied AI Environments[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 16277-16287.',
    '[38] Zhai G, Örnek E P, Wu S C, et al. CommonScenes: Generating Commonsense 3D Indoor Scenes with Scene Graph Diffusion[C]//Advances in Neural Information Processing Systems 36. 2023: 30026-30038.',
    '[39] Tang J, Nie Y, Markhasin L, et al. DiffuScene: Denoising Diffusion Models for Generative Indoor Scene Synthesis[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 20507-20518.',
    '[40] Ding L, Dong S, Huang Z, et al. Text-to-3D Generation with Bidirectional Diffusion Using Both 2D and 3D Priors[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 5115-5124.',
]

ref_overrides = {
    2: 'Tao F, Zhang H, Liu A, Nee A Y C. Digital Twin in Industry: State-of-the-Art[J]. IEEE Transactions on Industrial Informatics, 2019, 15(4): 2405-2415. DOI: 10.1109/TII.2018.2873186.',
    3: 'Tao F, Sui F, Liu A, et al. Digital twin-driven product design, manufacturing and service with big data[J]. The International Journal of Advanced Manufacturing Technology, 2018, 94(9-12): 3563-3576. DOI: 10.1007/s00170-017-0233-1.',
    4: 'Jones D, Snider C, Nassehi A, Yon J, Hicks B. Characterising the Digital Twin: A systematic literature review[J]. CIRP Journal of Manufacturing Science and Technology, 2020, 29: 36-52. DOI: 10.1016/j.cirpj.2020.02.002.',
    5: 'Pylianidis C, Osinga S, Athanasiadis I N. Introducing digital twins to agriculture[J]. Computers and Electronics in Agriculture, 2021, 184: 105942. DOI: 10.1016/j.compag.2020.105942.',
    6: 'Wolfert S, Ge L, Verdouw C, Bogaardt M J. Big Data in Smart Farming: A review[J]. Agricultural Systems, 2017, 153: 69-80. DOI: 10.1016/j.agsy.2017.01.023.',
    7: 'Liakos K G, Busato P, Moshou D, Pearson S, Bochtis D. Machine Learning in Agriculture: A Review[J]. Sensors, 2018, 18(8): 2674. DOI: 10.3390/s18082674.',
    8: 'Drury B, Fernandes R, Moura M F, de Andrade Lopes A. A survey of semantic web technology for agriculture[J]. Information Processing in Agriculture, 2019, 6(4): 487-501. DOI: 10.1016/j.inpa.2019.02.001.',
    10: 'Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[EB/OL]. arXiv:2312.10997, 2023[2026-06-02]. https://arxiv.org/abs/2312.10997.',
    12: 'Schick T, Dwivedi-Yu J, Dessi R, et al. Toolformer: Language Models Can Teach Themselves to Use Tools[C]//Advances in Neural Information Processing Systems 36. 2023: 68539-68551. DOI: 10.52202/075280-2997.',
    13: 'Wang L, Ma C, Feng X, et al. A Survey on Large Language Model based Autonomous Agents[EB/OL]. arXiv:2308.11432, 2023[2026-06-02]. https://arxiv.org/abs/2308.11432.',
    14: 'Xi Z, Chen W, Guo X, et al. The Rise and Potential of Large Language Model Based Agents: A Survey[EB/OL]. arXiv:2309.07864, 2023[2026-06-02]. https://arxiv.org/abs/2309.07864.',
    15: 'Hogan A, Blomqvist E, Cochez M, et al. Knowledge Graphs[J]. ACM Computing Surveys, 2021, 54(4): 1-37. DOI: 10.1145/3447772.',
    16: 'Staab S, Studer R. Handbook on Ontologies[M]. Berlin: Springer, 2009. DOI: 10.1007/978-3-540-92673-3.',
    17: 'Berners-Lee T, Hendler J, Lassila O. The Semantic Web[J]. Scientific American, 2001, 284(5): 34-43. DOI: 10.1038/scientificamerican0501-34.',
    18: "d'Avila Garcez A, Lamb L C. Neurosymbolic AI: The 3rd wave[J]. Artificial Intelligence Review, 2023, 56(11): 12387-12406. DOI: 10.1007/s10462-023-10448-w.",
    19: 'Arrieta A B, Diaz-Rodriguez N, Del Ser J, et al. Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI[J]. Information Fusion, 2020, 58: 82-115. DOI: 10.1016/j.inffus.2019.12.012.',
    20: 'Amodei D, Olah C, Steinhardt J, et al. Concrete Problems in AI Safety[EB/OL]. arXiv:1606.06565, 2016[2026-06-02]. https://arxiv.org/abs/1606.06565.',
    21: 'Xiang J, Lv Z, Xu S, et al. Structured 3D Latents for Scalable and Versatile 3D Generation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2025. DOI: 10.1109/CVPR52734.2025.02000.',
    22: 'Verdouw C, Tekinerdogan B, Beulens A, Wolfert S. Digital twins in smart farming[J]. Agricultural Systems, 2021, 189: 103046. DOI: 10.1016/j.agsy.2020.103046.',
    23: 'Walter A, Finger R, Huber R, Buchmann N. Smart farming is key to developing sustainable agriculture[J]. Proceedings of the National Academy of Sciences, 2017, 114(24): 6148-6150. DOI: 10.1073/pnas.1707462114.',
    24: 'Kamilaris A, Prenafeta-Boldú F X. Deep learning in agriculture: A survey[J]. Computers and Electronics in Agriculture, 2018, 147: 70-90. DOI: 10.1016/j.compag.2018.02.016.',
    25: 'Jonquet C, Toulet A, Arnaud E, et al. AgroPortal: A vocabulary and ontology repository for agronomy[J]. Computers and Electronics in Agriculture, 2018, 144: 126-143. DOI: 10.1016/j.compag.2017.10.012.',
    26: 'Compton M, Barnaghi P, Bermudez L, et al. The SSN ontology of the W3C semantic sensor network incubator group[J]. Journal of Web Semantics, 2012, 17: 25-32. DOI: 10.1016/j.websem.2012.05.003.',
    27: 'Janowicz K, Haller A, Cox S J D, Le Phuoc D, Lefrançois M. SOSA: A lightweight ontology for sensors, observations, samples, and actuators[J]. Journal of Web Semantics, 2019, 56: 1-10. DOI: 10.1016/j.websem.2018.06.003.',
    28: "Park J S, O'Brien J, Cai C J, Morris M R, Liang P, Bernstein M S. Generative Agents: Interactive Simulacra of Human Behavior[C]//Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology. 2023: 1-22. DOI: 10.1145/3586183.3606763.",
    29: 'Shinn N, Cassano F, Gopinath A, Narasimhan K, Yao S. Reflexion: Language Agents with Verbal Reinforcement Learning[C]//Advances in Neural Information Processing Systems 36. 2023: 8634-8652. DOI: 10.52202/075280-0377.',
    30: 'Li G, Hammoud H, Itani H, Khizbullin D, Ghanem B. CAMEL: Communicative Agents for "Mind" Exploration of Large Language Model Society[C]//Advances in Neural Information Processing Systems 36. 2023: 51991-52008. DOI: 10.52202/075280-2264.',
    31: 'Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models[C]//Advances in Neural Information Processing Systems 35. 2022: 24824-24837. DOI: 10.52202/068431-1800.',
    32: 'Kojima T, Gu S S, Reid M, Matsuo Y, Iwasawa Y. Large Language Models Are Zero-Shot Reasoners[C]//Advances in Neural Information Processing Systems 35. 2022: 22199-22213. DOI: 10.52202/068431-1613.',
    33: 'Yao S, Yu D, Zhao J, et al. Tree of Thoughts: Deliberate Problem Solving with Large Language Models[C]//Advances in Neural Information Processing Systems 36. 2023: 11809-11822. DOI: 10.52202/075280-0517.',
    34: 'Lin C H, Gao J, Tang L, et al. Magic3D: High-Resolution Text-to-3D Content Creation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2023: 300-309. DOI: 10.1109/CVPR52729.2023.00037.',
    35: 'Höllein L, Cao A, Owens A, Johnson J, Nießner M. Text2Room: Extracting Textured 3D Meshes from 2D Text-to-Image Models[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2023: 7875-7886. DOI: 10.1109/ICCV51070.2023.00727.',
    36: 'Kerbl B, Kopanas G, Leimkühler T, Drettakis G. 3D Gaussian Splatting for Real-Time Radiance Field Rendering[J]. ACM Transactions on Graphics, 2023, 42(4): 139. DOI: 10.1145/3592433.',
    37: 'Yang Y, Sun F Y, Weihs L, et al. Holodeck: Language Guided Generation of 3D Embodied AI Environments[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 16277-16287. DOI: 10.1109/CVPR52733.2024.01536.',
    38: 'Zhai G, Örnek E P, Wu S C, et al. CommonScenes: Generating Commonsense 3D Indoor Scenes with Scene Graph Diffusion[C]//Advances in Neural Information Processing Systems 36. 2023: 30026-30038. DOI: 10.52202/075280-1307.',
    39: 'Tang J, Nie Y, Markhasin L, et al. DiffuScene: Denoising Diffusion Models for Generative Indoor Scene Synthesis[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 20507-20518. DOI: 10.1109/CVPR52733.2024.01938.',
    40: 'Ding L, Dong S, Huang Z, et al. Text-to-3D Generation with Bidirectional Diffusion Using Both 2D and 3D Priors[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024: 5115-5124. DOI: 10.1109/CVPR52733.2024.00489.',
}

old_refs = refs[:]
for old_num, body in ref_overrides.items():
    old_refs[old_num - 1] = f'[{old_num}] {body}'

refs = []
for new_num, old_num in enumerate(REFERENCE_ORDER, start=1):
    body = re.sub(r'^\[\d+\]\s*', '', old_refs[old_num - 1])
    refs.append(f'[{new_num}] {body}')

for ref in refs:
    add_ref(doc, ref)

# ============================================================
# SAVE
# ============================================================

doc.save(OUTPUT)
print(f"Paper saved to: {OUTPUT}")
print("Done!")
