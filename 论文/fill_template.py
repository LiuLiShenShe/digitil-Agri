#!/usr/bin/env python3
"""
Fill the user's markdown content into the 计算机研究与发展 DOCX template.
Preserves template formatting, uses image placeholders.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import re
import os

TEMPLATE_PATH = "/data/fj/数字孪生-paper-work/论文/20251126正文模版.docx"
OUTPUT_PATH = "/data/fj/数字孪生-paper-work/论文/KAFarmTwin-投稿稿.docx"
MD_PATH = "/data/fj/数字孪生-paper-work/计算机研究与发展专题投稿初稿.md"

with open(MD_PATH, 'r', encoding='utf-8') as f:
    md_text = f.read()

doc = Document(TEMPLATE_PATH)

def clear_paragraph(para):
    """Remove all runs from a paragraph."""
    for run in para.runs:
        run._element.getparent().remove(run._element)

# Stop sections - these are handled separately
STOP_SECTIONS = {'6 结论', '作者简介与照片占位', '参考文献', '## 作者简介与照片占位'}

def parse_markdown_sections(md_text):
    """Parse markdown into sections for the paper body. Stops at 结论/作者简介/参考文献."""
    lines = md_text.split('\n')
    sections = []
    current_section = {'heading': '', 'level': 0, 'content': []}
    in_body = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('## 1 引言'):
            in_body = True
            if current_section['content'] or current_section['heading']:
                sections.append(current_section)
            current_section = {'heading': '1  引言', 'level': 1, 'content': []}
            continue

        if not in_body:
            continue

        if stripped.startswith('## '):
            heading = stripped[3:].strip()
            # Stop at conclusion, author info, or references
            if heading in STOP_SECTIONS:
                if current_section['content'] or current_section['heading']:
                    sections.append(current_section)
                current_section = {'heading': '', 'level': 0, 'content': []}  # clear to avoid duplicate
                break
            if current_section['content'] or current_section['heading']:
                sections.append(current_section)
            current_section = {'heading': heading, 'level': 1, 'content': []}
        elif stripped.startswith('### '):
            heading = stripped[4:].strip()
            if heading in STOP_SECTIONS:
                if current_section['content'] or current_section['heading']:
                    sections.append(current_section)
                break
            if current_section['content'] or current_section['heading']:
                sections.append(current_section)
            current_section = {'heading': heading, 'level': 2, 'content': []}
        else:
            current_section['content'].append(line)

    if current_section['content'] or current_section['heading']:
        sections.append(current_section)

    return sections

def process_content(content_lines):
    """Process content lines into structured elements."""
    elements = []
    i = 0
    while i < len(content_lines):
        line = content_lines[i].strip()

        if not line:
            i += 1
            continue

        # Image placeholder
        if (line.startswith('【图') or line.startswith('【表')) and '占位' in line:
            typ = 'image_placeholder' if '图' in line[:4] else 'table_placeholder'
            text = line.strip('【】')
            elements.append({'type': typ, 'text': text})
            while i + 1 < len(content_lines) and content_lines[i+1].strip().startswith('建议'):
                i += 1
            i += 1
            continue

        # Image reference
        if line.startswith('![') and '.png' in line:
            elements.append({'type': 'image_reference', 'text': line})
            i += 1
            continue

        # Markdown table
        if line.startswith('|') and i+1 < len(content_lines) and '---' in content_lines[i+1]:
            table_lines = []
            while i < len(content_lines) and content_lines[i].strip().startswith('|'):
                table_lines.append(content_lines[i].strip())
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue

        # Bold table title (standalone)
        if line.startswith('**表') and line.endswith('**'):
            elements.append({'type': 'table_title', 'text': line.strip('*')})
            i += 1
            continue

        # LaTeX math
        if line.startswith('\\['):
            math_lines = [line]
            i += 1
            while i < len(content_lines) and not content_lines[i].strip().startswith('\\]'):
                math_lines.append(content_lines[i])
                i += 1
            if i < len(content_lines):
                math_lines.append(content_lines[i])
                i += 1
            elements.append({'type': 'math', 'lines': math_lines})
            continue

        # Regular paragraph
        elements.append({'type': 'paragraph', 'text': line})
        i += 1

    return elements

# ============================================================
# STEP 1: Update front matter
# ============================================================

print("Step 1: Updating front matter...")

lines = md_text.split('\n')

chinese_title = "面向设施农业数字孪生的知识增强多智能体协作与可追溯场景构建方法"
english_title = "Knowledge-Augmented Multi-Agent Collaboration with Traceable Scene Construction for Protected Agriculture Digital Twins"

# Extract abstract/keywords
chinese_abstract = ""
english_abstract = ""
chinese_keywords = ""
english_keywords = ""

state = None
for line in lines:
    s = line.strip()
    if s == '## 摘要':
        state = 'cn_abs'
        continue
    if s == '## Abstract':
        state = 'en_abs'
        continue
    if s.startswith('**关键词**'):
        state = 'cn_kw'
        chinese_keywords = s.replace('**关键词**：', '').replace('**', '').strip()
        continue
    if s.startswith('**Keywords**'):
        state = None
        english_keywords = s.replace('**Keywords**:', '').replace('**', '').strip()
        continue
    if s.startswith('## 1 引言'):
        break

    if state == 'cn_abs' and s:
        chinese_abstract += s
    if state == 'en_abs' and s:
        english_abstract += s + ' '

english_abstract = english_abstract.strip()
chinese_abstract = chinese_abstract.strip()

print(f"  CN abstract: {len(chinese_abstract)} chars")
print(f"  EN abstract: {len(english_abstract)} chars")

# Apply front matter
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()

    if '基于分片区块链的车联网数据共享方案' in t and '题目三号' in t:
        clear_paragraph(p)
        r = p.add_run(chinese_title)
        r.bold = True
        r.font.size = Pt(16)
        print(f"  P{i}: Chinese title updated")

    if 'Internet of Vehicles Data Sharing Scheme via Blockchain Sharding' in t:
        clear_paragraph(p)
        r = p.add_run(english_title)
        r.font.size = Pt(14)
        print(f"  P{i}: English title updated")

    if t.startswith('Abstract ') and len(t) > 50:
        clear_paragraph(p)
        r = p.add_run('Abstract ')
        r.bold = True
        r = p.add_run(english_abstract)
        print(f"  P{i}: English abstract updated")

    if t.startswith('Key words') and 'Internet of vehicles' in t:
        clear_paragraph(p)
        r = p.add_run('Key words  ')
        r.bold = True
        r = p.add_run(english_keywords)
        print(f"  P{i}: English keywords updated")

    if (t.startswith('摘要') or t.startswith('摘要 ')) and '高效安全的数据共享' in t:
        clear_paragraph(p)
        r = p.add_run('摘  要  ')
        r.bold = True
        r = p.add_run(chinese_abstract)
        print(f"  P{i}: Chinese abstract updated")

    if t.startswith('关键词') and '车联网' in t:
        clear_paragraph(p)
        p.style = doc.styles['关键词']
        r = p.add_run('关键词 ')
        r.bold = True
        r = p.add_run(chinese_keywords)
        print(f"  P{i}: Chinese keywords updated")

# ============================================================
# STEP 2: Remove old body paragraphs
# ============================================================

print("\nStep 2: Removing demo body content...")

body_start_idx = None
conclusion_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '正文五号宋体':
        body_start_idx = i
    if t.startswith('6结论') or t.startswith('6 结论'):
        conclusion_idx = i

print(f"  Body start: P{body_start_idx}, Conclusion: P{conclusion_idx}")

# Remove paragraphs between body_start+1 and conclusion_idx-1
paras_to_remove = list(range(body_start_idx + 1, conclusion_idx))
for idx in reversed(paras_to_remove):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

print(f"  Removed {len(paras_to_remove)} paragraphs")

# ============================================================
# STEP 3: Build and insert body content
# ============================================================

print("\nStep 3: Building body content...")

sections = parse_markdown_sections(md_text)
print(f"  Found {len(sections)} sections")

# Find anchor points
anchor_para = None
conclusion_para = None
for p in doc.paragraphs:
    t = p.text.strip()
    if t == '正文五号宋体':
        anchor_para = p
    if t.startswith('6结论') or t.startswith('6 结论'):
        conclusion_para = p

# Build elements
def make_heading_elem(doc, text, level):
    """Create a heading paragraph element using proper style references."""
    p_elem = doc._element.makeelement(qn('w:p'), {})
    pPr = p_elem.makeelement(qn('w:pPr'), {})
    p_elem.append(pPr)

    if level == 1:
        # 一级标题: 小四黑体 12pt bold, use Heading 2 style
        pStyle = pPr.makeelement(qn('w:pStyle'), {})
        pStyle.set(qn('w:val'), 'Heading 2')
        pPr.append(pStyle)
        # Copy numbering + spacing from template's Heading 2
        numPr = pPr.makeelement(qn('w:numPr'), {})
        ilvl = numPr.makeelement(qn('w:ilvl'), {})
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numId = numPr.makeelement(qn('w:numId'), {})
        numId.set(qn('w:val'), '0')
        numPr.append(numId)
        pPr.append(numPr)

        r = p_elem.makeelement(qn('w:r'), {})
        rPr = r.makeelement(qn('w:rPr'), {})
        rPr_b = rPr.makeelement(qn('w:b'), {})
        rPr.append(rPr_b)
        rPr_sz = rPr.makeelement(qn('w:sz'), {})
        rPr_sz.set(qn('w:val'), '24')  # 12pt
        rPr.append(rPr_sz)
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rFonts.set(qn('w:eastAsia'), '黑体')
        rPr.append(rFonts)
        r.append(rPr)
        t_elem = r.makeelement(qn('w:t'), {})
        t_elem.text = text
        t_elem.set(qn('xml:space'), 'preserve')
        r.append(t_elem)
        p_elem.append(r)
    else:
        # 二级标题: 五号黑体 10.5pt bold
        pStyle = pPr.makeelement(qn('w:pStyle'), {})
        pStyle.set(qn('w:val'), 'Normal')
        pPr.append(pStyle)
        # Add left indent for alignment with body text
        ind = pPr.makeelement(qn('w:ind'), {})
        ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)

        r = p_elem.makeelement(qn('w:r'), {})
        rPr = r.makeelement(qn('w:rPr'), {})
        rPr_b = rPr.makeelement(qn('w:b'), {})
        rPr.append(rPr_b)
        rPr_sz = rPr.makeelement(qn('w:sz'), {})
        rPr_sz.set(qn('w:val'), '21')  # 10.5pt
        rPr.append(rPr_sz)
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rFonts.set(qn('w:eastAsia'), '黑体')
        rPr.append(rFonts)
        r.append(rPr)
        t_elem = r.makeelement(qn('w:t'), {})
        t_elem.text = text
        t_elem.set(qn('xml:space'), 'preserve')
        r.append(t_elem)
        p_elem.append(r)

    return p_elem

def make_body_elem(doc, text):
    """Create a body text paragraph element (宋体五号, first-line indent)."""
    p_elem = doc._element.makeelement(qn('w:p'), {})
    pPr = p_elem.makeelement(qn('w:pPr'), {})
    pStyle = pPr.makeelement(qn('w:pStyle'), {})
    pStyle.set(qn('w:val'), 'Body Text Indent')
    pPr.append(pStyle)
    p_elem.append(pPr)

    r = p_elem.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rPr_sz = rPr.makeelement(qn('w:sz'), {})
    rPr_sz.set(qn('w:val'), '21')
    rPr.append(rPr_sz)
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    r.append(rPr)
    t_elem = r.makeelement(qn('w:t'), {})
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r.append(t_elem)
    p_elem.append(r)

    return p_elem

def make_placeholder_elem(doc, text, size='18'):
    """Create a centered gray placeholder paragraph."""
    p_elem = doc._element.makeelement(qn('w:p'), {})
    pPr = p_elem.makeelement(qn('w:pPr'), {})
    jc = pPr.makeelement(qn('w:jc'), {})
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p_elem.append(pPr)

    r = p_elem.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rPr_sz = rPr.makeelement(qn('w:sz'), {})
    rPr_sz.set(qn('w:val'), size)
    rPr.append(rPr_sz)
    color = rPr.makeelement(qn('w:color'), {})
    color.set(qn('w:val'), '808080')
    rPr.append(color)
    i_elem = rPr.makeelement(qn('w:i'), {})
    rPr.append(i_elem)
    r.append(rPr)
    t_elem = r.makeelement(qn('w:t'), {})
    t_elem.text = f'【{text}】'
    t_elem.set(qn('xml:space'), 'preserve')
    r.append(t_elem)
    p_elem.append(r)

    return p_elem

def make_centered_text_elem(doc, text, bold=False, size='18'):
    """Create a centered text paragraph."""
    p_elem = doc._element.makeelement(qn('w:p'), {})
    pPr = p_elem.makeelement(qn('w:pPr'), {})
    jc = pPr.makeelement(qn('w:jc'), {})
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p_elem.append(pPr)

    r = p_elem.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    if bold:
        rPr_b = rPr.makeelement(qn('w:b'), {})
        rPr.append(rPr_b)
    rPr_sz = rPr.makeelement(qn('w:sz'), {})
    rPr_sz.set(qn('w:val'), size)
    rPr.append(rPr_sz)
    r.append(rPr)
    t_elem = r.makeelement(qn('w:t'), {})
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r.append(t_elem)
    p_elem.append(r)

    return p_elem

def make_table_elem(doc, table_lines):
    """Create a table element from markdown table lines."""
    def parse_row(line):
        return [c.strip() for c in line.strip('|').split('|')]

    header = parse_row(table_lines[0])
    rows = [parse_row(l) for l in table_lines[2:]]
    num_cols = len(header)
    col_width = 9000 // max(num_cols, 1)

    tbl = doc._element.makeelement(qn('w:tbl'), {})

    # Table properties
    tblPr = tbl.makeelement(qn('w:tblPr'), {})
    tblW = tblPr.makeelement(qn('w:tblW'), {})
    tblW.set(qn('w:w'), '9000')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblBorders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.makeelement(qn(f'w:{border_name}'), {})
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Grid
    tblGrid = tbl.makeelement(qn('w:tblGrid'), {})
    for c in range(num_cols):
        gridCol = tblGrid.makeelement(qn('w:gridCol'), {})
        gridCol.set(qn('w:w'), str(col_width))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # Rows
    all_rows = [header] + rows
    is_header = True
    for row_cells in all_rows:
        tr = tbl.makeelement(qn('w:tr'), {})
        for cell_text in row_cells:
            tc = tr.makeelement(qn('w:tc'), {})
            tcPr = tc.makeelement(qn('w:tcPr'), {})
            tcW = tcPr.makeelement(qn('w:tcW'), {})
            tcW.set(qn('w:w'), str(col_width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)

            p = tc.makeelement(qn('w:p'), {})
            pPr = p.makeelement(qn('w:pPr'), {})
            jc = pPr.makeelement(qn('w:jc'), {})
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            p.append(pPr)

            r = p.makeelement(qn('w:r'), {})
            rPr = r.makeelement(qn('w:rPr'), {})
            if is_header:
                rPr_b = rPr.makeelement(qn('w:b'), {})
                rPr.append(rPr_b)
            rPr_sz = rPr.makeelement(qn('w:sz'), {})
            rPr_sz.set(qn('w:val'), '15')
            rPr.append(rPr_sz)
            r.append(rPr)
            t_elem = r.makeelement(qn('w:t'), {})
            t_elem.text = cell_text
            t_elem.set(qn('xml:space'), 'preserve')
            r.append(t_elem)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
        is_header = False

    return tbl

def make_math_elem(doc, math_lines):
    """Create a math equation placeholder element."""
    math_text = ' '.join([l.strip() for l in math_lines
                          if l.strip() not in ('\\[', '\\]', '')])
    if len(math_text) > 200:
        math_text = math_text[:197] + '...'

    p_elem = doc._element.makeelement(qn('w:p'), {})
    pPr = p_elem.makeelement(qn('w:pPr'), {})
    jc = pPr.makeelement(qn('w:jc'), {})
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p_elem.append(pPr)

    r = p_elem.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rPr_sz = rPr.makeelement(qn('w:sz'), {})
    rPr_sz.set(qn('w:val'), '21')
    rPr.append(rPr_sz)
    r.append(rPr)
    t_elem = r.makeelement(qn('w:t'), {})
    t_elem.text = f'[公式待录入: {math_text}]'
    t_elem.set(qn('xml:space'), 'preserve')
    r.append(t_elem)
    p_elem.append(r)

    return p_elem

# Process all sections and build elements
new_body_elements = []

for section in sections:
    heading = section['heading']
    level = section['level']
    content = section['content']

    print(f"  Processing: {heading}")

    # Add section heading
    new_body_elements.append(make_heading_elem(doc, heading, level))

    # Process content
    elements = process_content(content)
    for elem in elements:
        if elem['type'] == 'paragraph':
            text = elem['text']
            if text.strip():
                new_body_elements.append(make_body_elem(doc, text))

        elif elem['type'] in ('image_placeholder', 'table_placeholder'):
            new_body_elements.append(make_placeholder_elem(doc, elem['text']))

        elif elem['type'] == 'image_reference':
            match = re.search(r'!\[(.*?)\]', elem['text'])
            alt = match.group(1) if match else 'Image'
            new_body_elements.append(make_placeholder_elem(doc, f'{alt} — 图片文件待插入'))

        elif elem['type'] == 'table_title':
            new_body_elements.append(make_centered_text_elem(doc, elem['text'], bold=True))

        elif elem['type'] == 'table':
            new_body_elements.append(make_table_elem(doc, elem['lines']))
            # Add blank line after table
            new_body_elements.append(make_body_elem(doc, ''))

        elif elem['type'] == 'math':
            new_body_elements.append(make_math_elem(doc, elem['lines']))

# Insert all body elements after anchor_para, before conclusion_para
if anchor_para and conclusion_para:
    current = anchor_para._element
    for elem in new_body_elements:
        current.addnext(elem)
        current = elem
    print(f"  Inserted {len(new_body_elements)} body elements")

# ============================================================
# STEP 4: Update conclusion
# ============================================================

print("\nStep 4: Updating conclusion...")

# Extract conclusion from markdown
conclusion_text = ""
in_conclusion = False
for line in lines:
    s = line.strip()
    if s.startswith('## 6 结论'):
        in_conclusion = True
        continue
    if in_conclusion:
        if s.startswith('## ') or s.startswith('# '):
            break
        if s:
            conclusion_text += s

# Update template's conclusion heading and content
# Look for the original template conclusion (heading "6结论" + content paragraph)
conclusion_heading_found = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()

    # Update conclusion heading: "6结论" → "6  结论"
    if t.startswith('6结论') and not conclusion_heading_found:
        clear_paragraph(p)
        r = p.add_run('6  结论')
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        conclusion_heading_found = True
        print(f"  P{i}: Conclusion heading updated")
        continue

    # Update conclusion content (the paragraph right after the heading)
    if conclusion_heading_found and ('本文提出了一种基于机器学习分片的车联网区块链数据共享方案' in t):
        clear_paragraph(p)
        p.style = doc.styles['Body Text Indent']
        r = p.add_run(conclusion_text)
        r.font.size = Pt(10.5)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        print(f"  P{i}: Conclusion content updated")
        conclusion_heading_found = False
        break

# Also clean up any markdown "6 结论" heading that might have been inserted
# and any old demo text that might remain
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # If there's an extra "6 结论" from markdown (not the template one)
    if t == '6 结论' and p.style.name == 'Normal':
        # Check if this is an extra one (not the one we just updated)
        runs_have_bold = any(r.bold for r in p.runs if r.bold)
        if not runs_have_bold:
            # This is the markdown heading - remove it and the next content paragraph
            p._element.getparent().remove(p._element)
            # The content paragraph right after it would now shift
            print(f"  P{i}: Removed duplicate '6 结论' from markdown body")

# ============================================================
# STEP 5: Update references
# ============================================================

print("\nStep 5: Updating references...")

# Extract references from markdown
refs = []
in_refs = False
for line in lines:
    s = line.strip()
    if s == '## 参考文献':
        in_refs = True
        continue
    if in_refs:
        if s.startswith('## ') or s.startswith('# '):
            break
        if s.startswith('[') and ']' in s[:6]:
            refs.append(s)

print(f"  Found {len(refs)} references")

# Find reference section boundaries
ref_heading_idx = None
author_info_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '参 考 文 献' in t:
        ref_heading_idx = i
    if '作者介绍小五号' in t:
        author_info_idx = i

# Remove old reference paragraphs
if ref_heading_idx and author_info_idx:
    ref_remove = list(range(ref_heading_idx + 1, author_info_idx))
    for idx in reversed(ref_remove):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    print(f"  Removed {len(ref_remove)} old reference paragraphs")

# Find ref heading again and insert new refs
ref_heading = None
for p in doc.paragraphs:
    if '参 考 文 献' in p.text.strip():
        ref_heading = p
        break

if ref_heading:
    current = ref_heading._element
    for ref_text in refs:
        p_elem = doc._element.makeelement(qn('w:p'), {})
        r = p_elem.makeelement(qn('w:r'), {})
        rPr = r.makeelement(qn('w:rPr'), {})
        rPr_sz = rPr.makeelement(qn('w:sz'), {})
        rPr_sz.set(qn('w:val'), '15')
        rPr.append(rPr_sz)
        r.append(rPr)
        t_elem = r.makeelement(qn('w:t'), {})
        t_elem.text = ref_text
        t_elem.set(qn('xml:space'), 'preserve')
        r.append(t_elem)
        p_elem.append(r)
        current.addnext(p_elem)
        current = p_elem
    print(f"  Inserted {len(refs)} new references")

# ============================================================
# STEP 6: Update author contribution and bios
# ============================================================

print("\nStep 6: Updating author info...")

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()

    if '作者贡献声明' in t and '陈骁' in t:
        clear_paragraph(p)
        r = p.add_run('作者贡献声明：')
        r.bold = True
        r = p.add_run('作者待填。')
        print(f"  P{i}: Author contribution placeholder")

    if t.startswith('Chen X*') and 'blockchain' in t.lower():
        clear_paragraph(p)
        r = p.add_run('[Author name], born in [year]. [degree]. [Research interests].')
        r.font.size = Pt(9)
        r.bold = True
        r2 = p.add_run(' (To be completed)')
        r2.font.size = Pt(9)
        print(f"  P{i}: Author bio EN 1 placeholder")

    if t.startswith('陈*') and '区块链' in t:
        clear_paragraph(p)
        r = p.add_run('[作者姓名]，[年份]年生。[学历]。[研究方向]。（待补全）')
        r.font.size = Pt(9)
        r.bold = True
        print(f"  P{i}: Author bio CN 1 placeholder")

    if t.startswith('Huang *hong') and 'blockchain' in t.lower():
        clear_paragraph(p)
        r = p.add_run('[Author name], born in [year]. [degree]. [Research interests].')
        r.font.size = Pt(9)
        r.bold = True
        r2 = p.add_run(' (To be completed)')
        r2.font.size = Pt(9)
        print(f"  P{i}: Author bio EN 2 placeholder")

    if t.startswith('黄*鸿') and '区块链' in t:
        clear_paragraph(p)
        r = p.add_run('[作者姓名]，[年份]年生。[学历]。[研究方向]。（待补全）')
        r.font.size = Pt(9)
        r.bold = True
        print(f"  P{i}: Author bio CN 2 placeholder")

# ============================================================
# STEP 7: Clean up orphaned instructional text and old template tables
# ============================================================

print("\nStep 7: Cleaning up...")

# Remove template-specific instruction paragraphs
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '三线表，表题为中英文，小五号黑体，表的内容尽量用中文，除变量、名称缩写外。':
        clear_paragraph(p)
        print(f"  P{i}: Cleared table formatting note")
    if t.startswith('中英文图题用小五号'):
        clear_paragraph(p)
        print(f"  P{i}: Cleared figure formatting note")
    if t.startswith('图例尽量用不同图案'):
        clear_paragraph(p)
        print(f"  P{i}: Cleared figure legend note")
    if '公式用office编辑器' in t:
        clear_paragraph(p)
        print(f"  P{i}: Cleared formula editor note")

# Remove old template demo tables (cloud server config)
tables_to_remove = []
for ti, table in enumerate(doc.tables):
    hdr_texts = [c.text.strip() for c in table.rows[0].cells]
    if hdr_texts == ['名称', '配置环境']:
        tables_to_remove.append(ti)
        print(f"  Found old template table {ti}: {hdr_texts}")

for ti in reversed(tables_to_remove):
    tbl = doc.tables[ti]
    tbl._element.getparent().remove(tbl._element)
    print(f"  Removed old template table {ti}")

# ============================================================
# Save
# ============================================================

print(f"\nSaving to {OUTPUT_PATH}...")
doc.save(OUTPUT_PATH)
print("Done!")
print(f"Output: {OUTPUT_PATH}")
