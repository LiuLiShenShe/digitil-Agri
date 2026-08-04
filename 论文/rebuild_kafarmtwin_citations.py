# -*- coding: utf-8 -*-
"""Rebuild KAFarmTwin inline citations and bibliography.

This script keeps the original DOCX unchanged. It creates a revised copy with
old numeric citations removed, new first-appearance citations inserted, and the
old bibliography replaced by a verified recent-reference list.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = ROOT / "KAFarmTwin_公式Office可编辑版.docx"
OUTPUT_DOCX = ROOT / "KAFarmTwin_清除旧引文并重排参考文献.docx"
MAPPING_MD = ROOT / "KAFarmTwin_引用映射与参考文献.md"

CITE_PATTERN = re.compile(r"\[[0-9,\-\u2013\u2014\s]+\]")


REFERENCES: list[str] = [
    "[1] Semeraro C, Lezoche M, Panetto H, et al. Digital twin paradigm: a systematic literature review[J]. Computers in Industry, 2021, 130: 103469. DOI: 10.1016/j.compind.2021.103469.",
    "[2] Botin-Sanabria D M, Mihaita A S, Peimbert-Garcia R E, et al. Digital twin technology challenges and applications: a comprehensive review[J]. Remote Sensing, 2022, 14(6): 1335. DOI: 10.3390/rs14061335.",
    "[3] Dihan M S, Akash A I, Tasneem Z, et al. Digital twin: data exploration, architecture, implementation and future[J]. Heliyon, 2024, 10(5): e26503. DOI: 10.1016/j.heliyon.2024.e26503.",
    "[4] Pylianidis C, Osinga S, Athanasiadis I N. Introducing digital twins to agriculture[J]. Computers and Electronics in Agriculture, 2021, 184: 105942. DOI: 10.1016/j.compag.2020.105942.",
    "[5] Verdouw C, Tekinerdogan B, Beulens A, et al. Digital twins in smart farming[J]. Agricultural Systems, 2021, 189: 103046. DOI: 10.1016/j.agsy.2020.103046.",
    "[6] Nasirahmadi A, Hensel O. Toward the next generation of digitalization in agriculture based on digital twin paradigm[J]. Sensors, 2022, 22(2): 498. DOI: 10.3390/s22020498.",
    "[7] Cesco S, Sambo P, Borin M, et al. Smart agriculture and digital twins: applications and challenges in a vision of sustainability[J]. European Journal of Agronomy, 2023, 146: 126809. DOI: 10.1016/j.eja.2023.126809.",
    "[8] Peladarinos N, Piromalis D, Cheimaras V, et al. Enhancing smart agriculture by implementing digital twins: a comprehensive review[J]. Sensors, 2023, 23(16): 7128. DOI: 10.3390/s23167128.",
    "[9] Escriva-Gelonch M, Liang Shu, van Schalkwyk P, et al. Digital twins in agriculture: orchestration and applications[J]. Journal of Agricultural and Food Chemistry, 2024, 72(19): 10737-10752. DOI: 10.1021/acs.jafc.4c01934.",
    "[10] Quy V K, Hau N V, Anh D V, et al. IoT-enabled smart agriculture: architecture, applications, and challenges[J]. Applied Sciences, 2022, 12(7): 3396. DOI: 10.3390/app12073396.",
    "[11] Finger R. Digital innovations for sustainable and resilient agricultural systems[J]. European Review of Agricultural Economics, 2023, 50(4): 1277-1309. DOI: 10.1093/erae/jbad021.",
    "[12] Dara R, Hazrati Fard S M, Kaur J. Recommendations for ethical and responsible use of artificial intelligence in digital agriculture[J]. Frontiers in Artificial Intelligence, 2022, 5: 884192. DOI: 10.3389/frai.2022.884192.",
    "[13] Wang Lei, Ma Chen, Feng Xueyang, et al. A survey on large language model based autonomous agents[J]. Frontiers of Computer Science, 2024, 18(6): 186345. DOI: 10.1007/s11704-024-40231-1.",
    "[14] Li Xinyi, Wang Sai, Zeng Siqi, et al. A survey on LLM-based multi-agent systems: workflow, infrastructure, and challenges[J]. Vicinagearth, 2024, 1(1): 9. DOI: 10.1007/s44336-024-00009-2.",
    "[15] Schick T, Dwivedi-Yu J, Dessi R, et al. Toolformer: language models can teach themselves to use tools[C]//Advances in Neural Information Processing Systems 36. La Jolla: Neural Information Processing Systems Foundation, 2023: 68539-68551. DOI: 10.52202/075280-2997.",
    "[16] Park J S, O'Brien J, Cai C J, et al. Generative agents: interactive simulacra of human behavior[C]//Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology. New York: Association for Computing Machinery, 2023: 1-22. DOI: 10.1145/3586183.3606763.",
    "[17] Shinn N, Cassano F, Gopinath A, et al. Reflexion: language agents with verbal reinforcement learning[C]//Advances in Neural Information Processing Systems 36. La Jolla: Neural Information Processing Systems Foundation, 2023: 8634-8652. DOI: 10.52202/075280-0377.",
    "[18] Li Guohao, Hammoud H, Itani H, et al. CAMEL: communicative agents for \"mind\" exploration of large language model society[C]//Advances in Neural Information Processing Systems 36. La Jolla: Neural Information Processing Systems Foundation, 2023: 51991-52008. DOI: 10.52202/075280-2264.",
    "[19] Wei Jason, Wang Xuezhi, Schuurmans D, et al. Chain-of-thought prompting elicits reasoning in large language models[C]//Advances in Neural Information Processing Systems 35. La Jolla: Neural Information Processing Systems Foundation, 2022: 24824-24837. DOI: 10.52202/068431-1800.",
    "[20] Yao Shunyu, Yu Dian, Zhao Jeffrey, et al. Tree of thoughts: deliberate problem solving with large language models[C]//Advances in Neural Information Processing Systems 36. La Jolla: Neural Information Processing Systems Foundation, 2023: 11809-11822. DOI: 10.52202/075280-0517.",
    "[21] Zhan Qiusi, Liang Zhixiang, Ying Zifan, et al. InjecAgent: benchmarking indirect prompt injections in tool-integrated large language model agents[C]//Findings of the Association for Computational Linguistics ACL 2024. Stroudsburg: Association for Computational Linguistics, 2024: 10471-10506. DOI: 10.18653/v1/2024.findings-acl.624.",
    "[22] Zhu Feiyu, Simmons R. Bootstrapping cognitive agents with a large language model[C]//Proceedings of the AAAI Conference on Artificial Intelligence. Palo Alto: Association for the Advancement of Artificial Intelligence, 2024: 655-663. DOI: 10.1609/aaai.v38i1.27822.",
    "[23] Zhang Xinyu, Xu Huiyu, Ba Zhongjie, et al. PrivacyAsst: safeguarding user privacy in tool-using large language model agents[J]. IEEE Transactions on Dependable and Secure Computing, 2024, 21(6): 5242-5258. DOI: 10.1109/TDSC.2024.3372777.",
    "[24] Bran A M, Cox S, Schilter O, et al. Augmenting large language models with chemistry tools[J]. Nature Machine Intelligence, 2024, 6(5): 525-535. DOI: 10.1038/s42256-024-00832-8.",
    "[25] Jia Jingyi, Li Qinbin. AutoTool: efficient tool selection for large language model agents[C]//Proceedings of the AAAI Conference on Artificial Intelligence. Palo Alto: Association for the Advancement of Artificial Intelligence, 2026: 31265-31273. DOI: 10.1609/aaai.v40i37.40389.",
    "[26] Gao Yunfan, Xiong Yun, Gao Xinyu, et al. Retrieval-augmented generation for large language models: a survey[EB/OL]. arXiv:2312.10997, 2023[2026-06-03]. https://arxiv.org/abs/2312.10997. DOI: 10.48550/arXiv.2312.10997.",
    "[27] Procko T T, Ochoa O. Graph retrieval-augmented generation for large language models: a survey[C]//2024 Conference on AI, Science, Engineering, and Technology. Piscataway: IEEE, 2024: 166-169. DOI: 10.1109/AIxSET62544.2024.00030.",
    "[28] Ibrahim N, Aboulela S, Ibrahim A, et al. A survey on augmenting knowledge graphs (KGs) with large language models (LLMs): models, evaluation metrics, benchmarks, and challenges[J]. Discover Artificial Intelligence, 2024, 4(1): 76. DOI: 10.1007/s44163-024-00175-8.",
    "[29] Chen Guanyu, Song Tao, Wang Quanyu, et al. Knowledge graph and large language model integration with focus on educational applications: a survey[J]. Neurocomputing, 2025, 654: 131230. DOI: 10.1016/j.neucom.2025.131230.",
    "[30] d'Avila Garcez A, Lamb L C. Neurosymbolic AI: the 3rd wave[J]. Artificial Intelligence Review, 2023, 56(11): 12387-12406. DOI: 10.1007/s10462-023-10448-w.",
    "[31] Longo L, Brcic M, Cabitza F, et al. Explainable artificial intelligence (XAI) 2.0: a manifesto of open challenges and interdisciplinary research directions[J]. Information Fusion, 2024, 106: 102301. DOI: 10.1016/j.inffus.2024.102301.",
    "[32] Subagdja B, Shanthoshigaa D, Wang Zhaoxia, et al. Machine learning for refining knowledge graphs: a survey[J]. ACM Computing Surveys, 2024, 56(6): 1-38. DOI: 10.1145/3640313.",
    "[33] Amdouni E, Bouazzouni S, Jonquet C. O'FAIRe: ontology FAIRness evaluator in the AgroPortal semantic resource repository[C]//Lecture Notes in Computer Science. Cham: Springer International Publishing, 2022: 89-94. DOI: 10.1007/978-3-031-11609-4_17.",
    "[34] Chandra R, Agarwal S, Singh N. Semantic sensor network ontology based decision support system for forest fire management[J]. Ecological Informatics, 2022, 72: 101821. DOI: 10.1016/j.ecoinf.2022.101821.",
    "[35] Milli M, Milli M, Lakestani S, et al. Semantic-based anomaly detection in laboratory environments using SOSA/SSN sensor ontology frameworks[J]. Pamukkale University Journal of Engineering Sciences, 2023, 29(4): 357-369. DOI: 10.5505/pajes.2022.95595.",
    "[36] Lin Chen-Hsuan, Gao Jun, Tang Luming, et al. Magic3D: high-resolution text-to-3D content creation[C]//2023 IEEE/CVF Conference on Computer Vision and Pattern Recognition. Piscataway: IEEE, 2023: 300-309. DOI: 10.1109/CVPR52729.2023.00037.",
    "[37] Liu Qihao, Zhang Yi, Bai Song, et al. DIRECT-3D: learning direct text-to-3D generation on massive noisy 3D data[C]//2024 IEEE/CVF Conference on Computer Vision and Pattern Recognition. Piscataway: IEEE, 2024: 6881-6891. DOI: 10.1109/CVPR52733.2024.00657.",
    "[38] Tsalicoglou C, Manhardt F, Tonioni A, et al. TextMesh: generation of realistic 3D meshes from text prompts[C]//2024 International Conference on 3D Vision. Piscataway: IEEE, 2024: 1554-1563. DOI: 10.1109/3DV62453.2024.00154.",
    "[39] Yang Yue, Sun Fan-Yun, Weihs L, et al. Holodeck: language guided generation of 3D embodied AI environments[C]//2024 IEEE/CVF Conference on Computer Vision and Pattern Recognition. Piscataway: IEEE, 2024: 16277-16287. DOI: 10.1109/CVPR52733.2024.01536.",
    "[40] Xiang Jianfeng, Lv Zelong, Xu Sicheng, et al. Structured 3D latents for scalable and versatile 3D generation[C]//2025 IEEE/CVF Conference on Computer Vision and Pattern Recognition. Piscataway: IEEE, 2025: 21469-21480. DOI: 10.1109/CVPR52734.2025.02000.",
]


CITATION_MAP: dict[int, str] = {
    12: "[1-12]",
    13: "[4-12]",
    14: "[13-25]",
    24: "[1-12]",
    26: "[13-25]",
    28: "[26-33]",
    30: "[13,16-17,21,23,30-31]",
    32: "[1-12,13-35]",
    36: "[1-12,13-35]",
    50: "[32-35]",
    51: "[32-35]",
    57: "[26-35]",
    60: "[26-35]",
    65: "[13-25]",
    69: "[13-25]",
    77: "[36-40]",
    81: "[36-40]",
    83: "[36-40]",
    85: "[34-35]",
    91: "[34-35]",
    95: "[34-35]",
    98: "[21,23,30-31]",
    101: "[21,23,30-31]",
    106: "[16-17,21,23,30-31]",
    107: "[16-17,21,23,30-31]",
    111: "[16-17,21,23,30-31]",
    114: "[1-12,13-25]",
    117: "[1-12,13-25,36-40]",
    121: "[1-12,13-25,36-40]",
    136: "[13-31]",
    139: "[13-31]",
    143: "[1-3,13-31]",
    148: "[13-31]",
    190: "[1-40]",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def clean_old_citations(paragraph) -> None:
    for run in paragraph.runs:
        if run.text:
            run.text = CITE_PATTERN.sub("", run.text)


def append_citation(paragraph, citation: str) -> None:
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        paragraph.add_run(citation)
        return

    run = text_runs[-1]
    terminal = "。；;.!?！？"
    stripped = run.text.rstrip()
    trailing_spaces = run.text[len(stripped) :]
    if stripped and stripped[-1] in terminal:
        run.text = stripped[:-1] + citation + stripped[-1] + trailing_spaces
    else:
        run.text = run.text + citation


def rebuild_docx() -> None:
    doc = Document(SOURCE_DOCX)

    ref_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text in {"参 考 文 献", "参考文献"}:
            ref_idx = idx
            break

    if ref_idx is None:
        raise RuntimeError("Could not locate bibliography heading.")

    for idx, paragraph in enumerate(doc.paragraphs[:ref_idx]):
        clean_old_citations(paragraph)
        if idx in CITATION_MAP:
            append_citation(paragraph, CITATION_MAP[idx])

    for paragraph in list(doc.paragraphs[ref_idx:]):
        delete_paragraph(paragraph)

    heading = doc.add_paragraph("参 考 文 献")
    heading.alignment = 1
    for ref in REFERENCES:
        doc.add_paragraph(ref)

    doc.save(OUTPUT_DOCX)


def write_mapping() -> None:
    lines: list[str] = []
    lines.append("# KAFarmTwin 引用映射与参考文献")
    lines.append("")
    lines.append(f"- 源文件: `{SOURCE_DOCX}`")
    lines.append(f"- 输出文件: `{OUTPUT_DOCX}`")
    lines.append(f"- 参考文献数量: {len(REFERENCES)}")
    lines.append("- 检索/整理日期: 2026-06-03")
    lines.append("")
    lines.append("## 段落引用映射")
    lines.append("")

    doc = Document(SOURCE_DOCX)
    for idx in sorted(CITATION_MAP):
        text = CITE_PATTERN.sub("", doc.paragraphs[idx].text.strip())
        text = re.sub(r"\s+", " ", text)
        lines.append(f"- P{idx:03d} {CITATION_MAP[idx]}: {text[:180]}")

    lines.append("")
    lines.append("## 参考文献")
    lines.append("")
    lines.extend(REFERENCES)
    lines.append("")
    MAPPING_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    rebuild_docx()
    write_mapping()
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {MAPPING_MD}")


if __name__ == "__main__":
    main()
